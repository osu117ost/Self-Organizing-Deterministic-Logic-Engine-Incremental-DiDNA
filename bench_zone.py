#!/usr/bin/env python3
from __future__ import annotations

import argparse
import csv
import datetime as dt
import hashlib
import html
import json
import math
import os
import platform
import random
import re
import resource
import shutil
import statistics
import subprocess
import sys
import textwrap
import uuid
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Sequence, Tuple

import numpy as np
import yaml


RESULT_REQUIRED_FIELDS: Sequence[str] = (
    "dataset_id",
    "tier",
    "scale",
    "algorithm",
    "build_time_s",
    "qps",
    "lat_p50_ms",
    "lat_p95_ms",
    "lat_p99_ms",
    "rss_mb",
    "index_bytes",
    "bytes_per_insert",
    "recall_at_k",
    "update_latency_ms",
    "energy_proxy",
    "seed",
    "run_id",
    "error_bits_min_hamming",
    "error_bits_prefix_gap",
)

RESULT_EXTRA_FIELDS: Sequence[str] = (
    "timestamp_utc",
    "track",
    "framework",
    "operation",
    "status",
    "reason",
    "repetition",
    "hamming_audit_gap",
    "online_accuracy",
    "online_f1",
    "adaptation_lag_steps",
    "recall_at_k_id_overlap",
    "recall_at_k_tie_aware",
    "legal_review_required",
    "source_id",
    "requested_scale_points",
    "effective_db_size",
    "effective_query_size",
    "container_hash",
    "hardware_fingerprint",
)

RESULT_FIELDS: Sequence[str] = tuple(list(RESULT_REQUIRED_FIELDS) + list(RESULT_EXTRA_FIELDS))

SCALE_CHOICES = ("1M", "10M", "100M")
TIER_CHOICES = ("server", "desktop", "edge")
TRACK_CHOICES = ("core", "ann", "online")

OFFICIAL_SOURCE_PACK = [
    "https://mlcommons.org/benchmarks/",
    "https://docs.mlcommons.org/inference/index_gh/",
    "https://github.com/erikbern/ann-benchmarks",
    "https://ann-benchmarks.com/",
    "https://github.com/harsha-simhadri/big-ann-benchmarks",
    "https://big-ann-benchmarks.com/neurips23.html",
    "https://github.com/facebookresearch/faiss/wiki/Binary-hashing-index-benchmark",
    "https://github.com/facebookresearch/faiss",
    "https://github.com/nmslib/hnswlib",
    "https://github.com/nmslib/nmslib",
    "https://github.com/learnedsystems/SOSD",
    "https://github.com/brianfrankcooper/YCSB",
    "https://raw.githubusercontent.com/brianfrankcooper/YCSB/master/workloads/workloada",
    "https://raw.githubusercontent.com/brianfrankcooper/YCSB/master/workloads/workloadf",
    "https://github.com/facebook/rocksdb/wiki/Benchmarking-tools",
    "https://github.com/google/benchmark",
    "https://github.com/eembc/coremark",
    "https://github.com/lsils/benchmarks",
    "https://people.eecs.berkeley.edu/~alanmi/benchmarks/table_ex/",
    "https://people.eecs.berkeley.edu/~alanmi/benchmarks/iwls2005/",
    "https://moa.cms.waikato.ac.nz/",
    "https://riverml.xyz/latest/",
    "https://www.openml.org/benchmark",
    "https://archive.ics.uci.edu/",
]


EXTERNAL_BENCH_LIBRARY: Sequence[Dict[str, str]] = (
    {
        "category": "ANN / Similarity",
        "name": "ANN-Benchmarks",
        "scope": "Recall-latency-memory tradeoffs across ANN implementations",
        "url": "https://github.com/erikbern/ann-benchmarks",
    },
    {
        "category": "ANN / Similarity",
        "name": "ANN-Benchmarks website plots",
        "scope": "Recall/QPS, Recall/Build time, Recall/Index size, latency percentiles",
        "url": "https://ann-benchmarks.com/",
    },
    {
        "category": "ANN / Similarity",
        "name": "Big-ANN (NeurIPS track)",
        "scope": "Large-scale practical vector search; normalized hardware, recall/QPS leaderboards",
        "url": "https://big-ann-benchmarks.com/neurips23.html",
    },
    {
        "category": "ANN / Binary",
        "name": "Faiss Binary Hashing Benchmark",
        "scope": "256-bit binary vectors; recall, wall-clock, distance computations, random accesses",
        "url": "https://github.com/facebookresearch/faiss/wiki/Binary-hashing-index-benchmark",
    },
    {
        "category": "Tree / KV serving",
        "name": "YCSB",
        "scope": "Read/write/update/scan workloads with Zipfian and related distributions",
        "url": "https://github.com/brianfrankcooper/YCSB",
    },
    {
        "category": "Tree / KV serving",
        "name": "YCSB workload A",
        "scope": "50/50 read-update, Zipfian request distribution",
        "url": "https://raw.githubusercontent.com/brianfrankcooper/YCSB/master/workloads/workloada",
    },
    {
        "category": "Tree / KV serving",
        "name": "YCSB workload F",
        "scope": "50/50 read-read-modify-write, Zipfian request distribution",
        "url": "https://raw.githubusercontent.com/brianfrankcooper/YCSB/master/workloads/workloadf",
    },
    {
        "category": "Tree / Storage engine",
        "name": "RocksDB db_bench",
        "scope": "Official benchmark tool with fill/read/seek/mixed workloads",
        "url": "https://github.com/facebook/rocksdb/wiki/Benchmarking-tools",
    },
    {
        "category": "Tree / Sorted index",
        "name": "SOSD (Search on Sorted Data)",
        "scope": "Sorted-data lookup benchmark with many datasets and baseline index structures",
        "url": "https://github.com/learnedsystems/SOSD",
    },
    {
        "category": "Truth-table / Logic compression",
        "name": "EPFL Combinational Benchmark Suite",
        "scope": "Combinational logic suite for synthesis/optimization (arithmetic/control/MtM)",
        "url": "https://github.com/lsils/benchmarks",
    },
    {
        "category": "Truth-table / Logic compression",
        "name": "Berkeley multi-output PLA tables (table_ex)",
        "scope": "PLA table benchmark package used in cube-hashing extraction work",
        "url": "https://people.eecs.berkeley.edu/~alanmi/benchmarks/table_ex/",
    },
    {
        "category": "Truth-table / Logic compression",
        "name": "IWLS 2005 benchmark bundle",
        "scope": "Widely used logic synthesis benchmark set (AIG/Verilog archives)",
        "url": "https://people.eecs.berkeley.edu/~alanmi/benchmarks/iwls2005/",
    },
)


def deep_merge(base: Dict[str, Any], override: Dict[str, Any]) -> Dict[str, Any]:
    out = dict(base)
    for k, v in override.items():
        if isinstance(v, dict) and isinstance(out.get(k), dict):
            out[k] = deep_merge(out[k], v)
        else:
            out[k] = v
    return out


def default_config() -> Dict[str, Any]:
    return {
        "output_root": "bench_runs",
        "seed": 123456789,
        "licensing_policy": "internal_review",
        "report_language": "both",
        "strict_scale": False,
        "core_tool_bin": "builded/bench_core256_tool",
        "core_tool_src": "bench_core256_tool.c",
        "core_tool_compile": {
            "cc": "gcc",
            "cflags": [
                "-std=c23",
                "-O2",
                "-Wall",
                "-Wextra",
                "-pedantic",
                "-Wno-unused-function",
            ],
            "sources": ["bench_core256_tool.c", "core.c", "core256.c", "core256_mv.c"],
            "ldflags": ["-pthread", "-ldl", "-lm"],
            "output": "builded/bench_core256_tool",
        },
        "global": {
            "repetitions": 5,
            "top_k": 10,
            "hamming_audit_every": 10,
            "hamming_audit_max_queries": 128,
            "legacy_compat_enabled": True,
        },
        "scale_profiles": {
            "1M": {
                "requested_points": 1_000_000,
                "core_samples": 200_000,
                "ann_db": 200_000,
                "ann_queries": 2_000,
                "online_events": 40_000,
            },
            "10M": {
                "requested_points": 10_000_000,
                "core_samples": 800_000,
                "ann_db": 800_000,
                "ann_queries": 5_000,
                "online_events": 120_000,
            },
            "100M": {
                "requested_points": 100_000_000,
                "core_samples": 2_000_000,
                "ann_db": 2_000_000,
                "ann_queries": 10_000,
                "online_events": 250_000,
            },
        },
        "tiers": {
            "server": {"notes": "high core count, large RAM"},
            "desktop": {"notes": "workstation class"},
            "edge": {"notes": "resource-constrained profile"},
        },
        "tracks": {
            "core": {
                "enabled": True,
                "framework": "google_benchmark_compat+native",
                "workloads": ["random", "clustered", "adversarial", "replay"],
            },
            "ann": {
                "enabled": True,
                "framework": "ann-benchmarks+big-ann+native",
                "datasets": [
                    {
                        "id": "synthetic_binary256",
                        "source": "synthetic",
                        "kind": "binary256_clustered",
                    },
                    {
                        "id": "openml:adult",
                        "source": "openml",
                        "name": "adult",
                    },
                ],
                "baselines": ["faiss", "hnswlib", "nmslib"],
                "max_ground_truth_queries": 128,
            },
            "online": {
                "enabled": True,
                "framework": "delayed_preq+drift+core256_similarity",
                "delay": 64,
                "drift_interval": 5000,
                "hamming_sample_rate": 10,
                "buffer_cap": 65536,
            },
        },
    }


def load_config(path: Path) -> Dict[str, Any]:
    cfg = default_config()
    if path.exists():
        user_cfg = yaml.safe_load(path.read_text(encoding="utf-8")) or {}
        if not isinstance(user_cfg, dict):
            raise ValueError(f"Config must be a mapping: {path}")
        cfg = deep_merge(cfg, user_cfg)
    return cfg


def hash_jsonable(value: Any) -> str:
    payload = json.dumps(value, sort_keys=True, separators=(",", ":"), ensure_ascii=False).encode("utf-8")
    return hashlib.sha256(payload).hexdigest()


def utc_now() -> str:
    return dt.datetime.now(dt.timezone.utc).replace(microsecond=0).isoformat().replace("+00:00", "Z")


def safe_run(cmd: Sequence[str], cwd: Optional[Path] = None) -> subprocess.CompletedProcess[str]:
    return subprocess.run(
        list(cmd),
        cwd=str(cwd) if cwd else None,
        text=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        check=False,
    )


def parse_tool_version(cmd: Sequence[str]) -> str:
    cp = safe_run(cmd)
    if cp.returncode == 0 and cp.stdout.strip():
        return cp.stdout.strip().splitlines()[0]
    if cp.stderr.strip():
        return cp.stderr.strip().splitlines()[0]
    return "unknown"


def detect_git_commit(repo_root: Path) -> Optional[str]:
    cp = safe_run(["git", "rev-parse", "HEAD"], cwd=repo_root)
    if cp.returncode == 0:
        v = cp.stdout.strip()
        if v:
            return v
    return None


def detect_container_hash() -> str:
    env_candidates = [
        os.environ.get("CONTAINER_ID"),
        os.environ.get("HOSTNAME"),
    ]
    for c in env_candidates:
        if c and len(c) >= 12:
            return c[:64]

    cgroup = Path("/proc/self/cgroup")
    if cgroup.exists():
        data = cgroup.read_text(encoding="utf-8", errors="ignore")
        m = re.findall(r"[0-9a-f]{12,64}", data)
        if m:
            return m[-1]
    return "host-process"


def detect_cpu_model() -> str:
    cpuinfo = Path("/proc/cpuinfo")
    if cpuinfo.exists():
        for line in cpuinfo.read_text(encoding="utf-8", errors="ignore").splitlines():
            if line.lower().startswith("model name"):
                parts = line.split(":", 1)
                if len(parts) == 2:
                    return parts[1].strip()
    return platform.processor() or "unknown"


def detect_mem_total_mb() -> Optional[float]:
    meminfo = Path("/proc/meminfo")
    if meminfo.exists():
        for line in meminfo.read_text(encoding="utf-8", errors="ignore").splitlines():
            if line.startswith("MemTotal:"):
                parts = line.split()
                if len(parts) >= 2 and parts[1].isdigit():
                    return float(parts[1]) / 1024.0
    return None


def hardware_fingerprint() -> Dict[str, Any]:
    return {
        "platform": platform.platform(),
        "machine": platform.machine(),
        "cpu_model": detect_cpu_model(),
        "cpu_count": os.cpu_count(),
        "mem_total_mb": detect_mem_total_mb(),
        "python": sys.version.splitlines()[0],
        "gcc": parse_tool_version(["gcc", "--version"]),
    }


def rss_mb() -> float:
    # Linux returns KB, macOS returns bytes; normalize approximately.
    value = resource.getrusage(resource.RUSAGE_SELF).ru_maxrss
    if sys.platform == "darwin":
        return float(value) / (1024.0 * 1024.0)
    return float(value) / 1024.0


def percentile(values: Sequence[float], p: float) -> float:
    if not values:
        return 0.0
    xs = sorted(values)
    idx = int(math.floor((len(xs) - 1) * p))
    idx = max(0, min(len(xs) - 1, idx))
    return float(xs[idx])


def numeric_mean(values: Iterable[Optional[float]]) -> Optional[float]:
    xs = [float(v) for v in values if v is not None]
    if not xs:
        return None
    return float(sum(xs) / len(xs))


def to_float(value: Any) -> Optional[float]:
    if value is None:
        return None
    try:
        v = float(value)
    except Exception:
        return None
    if math.isnan(v) or math.isinf(v):
        return None
    return v


def coefficient_of_variation(values: Iterable[Optional[float]]) -> Optional[float]:
    xs = [float(v) for v in values if v is not None]
    if len(xs) < 2:
        return None
    mean_val = statistics.fmean(xs)
    if mean_val == 0.0:
        return None
    return float(statistics.pstdev(xs) / abs(mean_val))


def write_jsonl(path: Path, rows: Sequence[Dict[str, Any]]) -> None:
    with path.open("w", encoding="utf-8") as f:
        for row in rows:
            f.write(json.dumps(row, ensure_ascii=False, sort_keys=True) + "\n")


def ensure_result_shape(row: Dict[str, Any]) -> Dict[str, Any]:
    out = {k: row.get(k) for k in RESULT_FIELDS}
    for k in RESULT_FIELDS:
        if k not in out:
            out[k] = None
    return out


def build_summary_rows(rows: Sequence[Dict[str, Any]]) -> List[Dict[str, Any]]:
    key_fields = ["dataset_id", "tier", "scale", "algorithm", "track", "status"]
    numeric_fields = [
        "build_time_s",
        "qps",
        "lat_p50_ms",
        "lat_p95_ms",
        "lat_p99_ms",
        "rss_mb",
        "index_bytes",
        "bytes_per_insert",
        "recall_at_k",
        "recall_at_k_id_overlap",
        "recall_at_k_tie_aware",
        "update_latency_ms",
        "energy_proxy",
        "error_bits_min_hamming",
        "error_bits_prefix_gap",
        "hamming_audit_gap",
        "online_accuracy",
        "online_f1",
        "adaptation_lag_steps",
    ]
    grouped: Dict[Tuple[Any, ...], List[Dict[str, Any]]] = {}
    for row in rows:
        key = tuple(row.get(k) for k in key_fields)
        grouped.setdefault(key, []).append(row)

    output: List[Dict[str, Any]] = []
    for key, group_rows in grouped.items():
        out: Dict[str, Any] = {k: v for k, v in zip(key_fields, key)}
        out["n_runs"] = len(group_rows)
        for nf in numeric_fields:
            out[nf] = numeric_mean([r.get(nf) for r in group_rows])
        output.append(out)
    output.sort(key=lambda r: tuple(str(r.get(k)) for k in key_fields))
    return output


def write_summary_csv(path: Path, rows: Sequence[Dict[str, Any]]) -> None:
    summary_rows = build_summary_rows(rows)
    if not summary_rows:
        fieldnames = ["dataset_id", "tier", "scale", "algorithm", "track", "status", "n_runs"]
        with path.open("w", encoding="utf-8", newline="") as f:
            w = csv.DictWriter(f, fieldnames=fieldnames)
            w.writeheader()
        return
    with path.open("w", encoding="utf-8", newline="") as f:
        fieldnames = list(summary_rows[0].keys())
        w = csv.DictWriter(f, fieldnames=fieldnames)
        w.writeheader()
        for row in summary_rows:
            w.writerow(row)


def prepare_run_dir(output_root: Path, run_id: str) -> Path:
    run_dir = output_root / run_id
    run_dir.mkdir(parents=True, exist_ok=True)
    return run_dir


def compile_core_tool(repo_root: Path, cfg: Dict[str, Any]) -> None:
    c = cfg["core_tool_compile"]
    output = repo_root / c["output"]
    sources = [repo_root / s for s in c["sources"]]
    if output.exists():
        out_mtime = output.stat().st_mtime
        newer = any(s.exists() and s.stat().st_mtime > out_mtime for s in sources)
        if not newer:
            return

    cmd = [c.get("cc", "gcc")]
    cmd += list(c.get("cflags", []))
    cmd += [str(s) for s in sources]
    cmd += list(c.get("ldflags", []))
    cmd += ["-o", str(output)]
    cp = safe_run(cmd, cwd=repo_root)
    if cp.returncode != 0:
        raise RuntimeError(f"Failed to compile core tool:\nSTDOUT:\n{cp.stdout}\nSTDERR:\n{cp.stderr}")


def base_row(run_id: str, seed: int, track: str, tier: str, scale: str, dataset_id: str, algorithm: str) -> Dict[str, Any]:
    row = {k: None for k in RESULT_FIELDS}
    row["timestamp_utc"] = utc_now()
    row["run_id"] = run_id
    row["seed"] = seed
    row["track"] = track
    row["tier"] = tier
    row["scale"] = scale
    row["dataset_id"] = dataset_id
    row["algorithm"] = algorithm
    row["status"] = "ok"
    return row


def run_core_track(repo_root: Path,
                   cfg: Dict[str, Any],
                   run_id: str,
                   tier: str,
                   scale: str,
                   seed: int,
                   scale_profile: Dict[str, Any]) -> List[Dict[str, Any]]:
    rows: List[Dict[str, Any]] = []
    core_cfg = cfg["tracks"]["core"]
    if not core_cfg.get("enabled", True):
        return rows

    tool = repo_root / cfg["core_tool_bin"]
    repetitions = int(cfg["global"]["repetitions"])
    samples = int(scale_profile["core_samples"])
    workloads = list(core_cfg.get("workloads", []))

    for rep in range(repetitions):
        for workload in workloads:
            cmd = [
                str(tool),
                "core",
                "--samples",
                str(samples),
                "--seed",
                str((seed + rep * 977 + len(workload) * 13) & 0xFFFFFFFFFFFFFFFF),
                "--workload",
                workload,
            ]
            cp = safe_run(cmd, cwd=repo_root)
            if cp.returncode != 0:
                row = base_row(run_id, seed, "core", tier, scale, f"core:{workload}", "core256/tool")
                row["status"] = "failed"
                row["reason"] = cp.stderr.strip() or cp.stdout.strip() or f"exit={cp.returncode}"
                row["framework"] = core_cfg.get("framework")
                row["repetition"] = rep
                row["requested_scale_points"] = int(scale_profile["requested_points"])
                rows.append(ensure_result_shape(row))
                continue

            for line in cp.stdout.splitlines():
                s = line.strip()
                if not s.startswith("{"):
                    continue
                payload = json.loads(s)
                op = str(payload.get("operation", "unknown"))
                row = base_row(run_id, seed, "core", tier, scale, f"core:{workload}", f"core256/{op}")
                row["framework"] = core_cfg.get("framework")
                row["operation"] = op
                row["repetition"] = rep
                row["requested_scale_points"] = int(scale_profile["requested_points"])
                row["effective_db_size"] = samples
                row["effective_query_size"] = samples
                row["build_time_s"] = float(payload.get("total_time_s", 0.0))
                row["qps"] = float(payload.get("qps", 0.0))
                row["lat_p50_ms"] = float(payload.get("lat_p50_ms", 0.0))
                row["lat_p95_ms"] = float(payload.get("lat_p95_ms", 0.0))
                row["lat_p99_ms"] = float(payload.get("lat_p99_ms", 0.0))
                row["rss_mb"] = rss_mb()
                row["index_bytes"] = int(payload.get("index_bytes", 0))
                row["bytes_per_insert"] = float(payload.get("bytes_per_insert", 0.0))
                row["recall_at_k"] = None
                row["update_latency_ms"] = None
                row["energy_proxy"] = None
                pg = payload.get("mean_prefix_gap")
                row["error_bits_prefix_gap"] = float(pg) if isinstance(pg, (int, float)) and pg >= 0 else None
                row["error_bits_min_hamming"] = None
                row["hamming_audit_gap"] = None
                rows.append(ensure_result_shape(row))

    if cfg["global"].get("legacy_compat_enabled", True):
        rows.extend(run_legacy_compat_benches(repo_root, cfg, run_id, tier, scale, seed, scale_profile))
    return rows


def run_legacy_compat_benches(repo_root: Path,
                              cfg: Dict[str, Any],
                              run_id: str,
                              tier: str,
                              scale: str,
                              seed: int,
                              scale_profile: Dict[str, Any]) -> List[Dict[str, Any]]:
    rows: List[Dict[str, Any]] = []
    legacy_specs = [
        ("builded/bench_tree", [str(min(int(scale_profile["core_samples"]), 50000)), "50000"], r"\(([\d.]+)\s+Mops/s\)"),
        ("builded/bench_truth", ["--quick"], None),
    ]
    for rel_bin, args, mops_pattern in legacy_specs:
        bin_path = repo_root / rel_bin
        if not bin_path.exists():
            row = base_row(run_id, seed, "core", tier, scale, "core:legacy_compat", f"compat/{Path(rel_bin).name}")
            row["framework"] = "legacy_compat"
            row["status"] = "skipped"
            row["reason"] = "binary not found"
            rows.append(ensure_result_shape(row))
            continue
        cp = safe_run([str(bin_path)] + args, cwd=repo_root)
        row = base_row(run_id, seed, "core", tier, scale, "core:legacy_compat", f"compat/{Path(rel_bin).name}")
        row["framework"] = "legacy_compat"
        row["operation"] = "compatibility_layer"
        row["requested_scale_points"] = int(scale_profile["requested_points"])
        row["effective_db_size"] = int(scale_profile["core_samples"])
        row["effective_query_size"] = None
        if cp.returncode != 0:
            row["status"] = "failed"
            row["reason"] = cp.stderr.strip() or cp.stdout.strip() or f"exit={cp.returncode}"
            rows.append(ensure_result_shape(row))
            continue
        row["status"] = "ok"
        row["reason"] = "compat run completed"
        row["build_time_s"] = None
        row["qps"] = None
        if mops_pattern:
            m = re.search(mops_pattern, cp.stdout)
            if m:
                try:
                    row["qps"] = float(m.group(1)) * 1_000_000.0
                except Exception:
                    row["qps"] = None
        row["lat_p50_ms"] = None
        row["lat_p95_ms"] = None
        row["lat_p99_ms"] = None
        row["rss_mb"] = rss_mb()
        row["index_bytes"] = None
        row["bytes_per_insert"] = None
        row["recall_at_k"] = None
        row["update_latency_ms"] = None
        row["energy_proxy"] = None
        rows.append(ensure_result_shape(row))
    return rows


def pack_bits_to_bytes(bits: np.ndarray) -> np.ndarray:
    packed = np.packbits(bits.astype(np.uint8), axis=1, bitorder="big")
    if packed.shape[1] > 32:
        return packed[:, :32]
    if packed.shape[1] < 32:
        pad = np.zeros((packed.shape[0], 32 - packed.shape[1]), dtype=np.uint8)
        return np.concatenate([packed, pad], axis=1)
    return packed


def coerce_matrix_numeric(x: np.ndarray) -> np.ndarray:
    arr = np.asarray(x)
    if arr.ndim != 2:
        raise ValueError("expected 2D feature matrix")
    if np.issubdtype(arr.dtype, np.number):
        return np.asarray(arr, dtype=np.float32)

    out = np.empty(arr.shape, dtype=np.float32)
    for col_idx in range(arr.shape[1]):
        col = arr[:, col_idx]
        try:
            out[:, col_idx] = np.asarray(col, dtype=np.float32)
            continue
        except Exception:
            pass

        values = np.asarray(col).astype(str)
        encoded = np.empty(values.shape[0], dtype=np.float32)
        categories: Dict[str, float] = {}
        next_code = 1.0
        for i, raw in enumerate(values):
            token = raw.strip()
            if token == "" or token.lower() in {"nan", "none", "null", "na", "?"}:
                encoded[i] = 0.0
                continue
            try:
                encoded[i] = float(token)
                continue
            except Exception:
                pass
            if token not in categories:
                categories[token] = next_code
                next_code += 1.0
            encoded[i] = categories[token]
        out[:, col_idx] = encoded
    return out


def project_to_binary256(x: np.ndarray, seed: int) -> np.ndarray:
    x = coerce_matrix_numeric(x)
    x = np.nan_to_num(x, nan=0.0, posinf=0.0, neginf=0.0)
    x = x - x.mean(axis=0, keepdims=True)
    dim = x.shape[1]
    if dim >= 256:
        projected = x[:, :256]
    else:
        rng = np.random.default_rng(seed)
        proj = rng.normal(0.0, 1.0, size=(dim, 256)).astype(np.float32)
        projected = x @ proj
    bits = projected > 0.0
    return pack_bits_to_bytes(bits)


def synthetic_binary_dataset(db_size: int, query_size: int, seed: int, kind: str) -> Tuple[np.ndarray, np.ndarray]:
    rng = np.random.default_rng(seed)
    if kind == "binary256_clustered":
        centers = rng.integers(0, 256, size=(16, 32), dtype=np.uint8)
        db = np.empty((db_size, 32), dtype=np.uint8)
        for i in range(db_size):
            c = centers[rng.integers(0, len(centers))]
            v = c.copy()
            flips = int(rng.integers(1, 5))
            for _ in range(flips):
                bit = int(rng.integers(0, 256))
                byte_idx = bit >> 3
                bit_mask = 1 << (bit & 7)
                v[byte_idx] ^= np.uint8(bit_mask)
            db[i] = v
        query = db[rng.integers(0, db_size, size=query_size)].copy()
        return db, query
    db = rng.integers(0, 256, size=(db_size, 32), dtype=np.uint8)
    query = rng.integers(0, 256, size=(query_size, 32), dtype=np.uint8)
    return db, query


def load_dataset(dataset_cfg: Dict[str, Any], db_size: int, query_size: int, seed: int) -> Tuple[np.ndarray, np.ndarray, str]:
    source = dataset_cfg.get("source", "synthetic")
    dataset_id = str(dataset_cfg.get("id", "unknown"))
    if source == "synthetic":
        kind = str(dataset_cfg.get("kind", "binary256_random"))
        db, query = synthetic_binary_dataset(db_size, query_size, seed, kind)
        return db, query, dataset_id

    if source == "openml":
        try:
            from sklearn.datasets import fetch_openml  # type: ignore
        except Exception as exc:
            raise RuntimeError(f"openml source requires scikit-learn: {exc}") from exc
        data_id = dataset_cfg.get("data_id")
        name = dataset_cfg.get("name")
        if data_id is None and not name:
            raise ValueError(f"openml dataset requires data_id or name: {dataset_id}")
        fetch_kwargs: Dict[str, Any] = {"as_frame": False, "parser": "auto"}
        if data_id is not None:
            fetch_kwargs["data_id"] = int(data_id)
        else:
            fetch_kwargs["name"] = str(name)
        try:
            ds = fetch_openml(**fetch_kwargs)
        except Exception as exc:
            # Some OpenML parsers require optional deps (e.g., pandas).
            # Fall back to liac-arff parser for reproducible minimal environments.
            fetch_kwargs["parser"] = "liac-arff"
            try:
                ds = fetch_openml(**fetch_kwargs)
            except Exception as fallback_exc:
                raise RuntimeError(
                    f"openml fetch failed for {dataset_id}: {exc}; fallback parser failed: {fallback_exc}"
                ) from fallback_exc
        x = np.asarray(ds.data)
        if x.shape[0] < db_size + query_size:
            reps = int(math.ceil((db_size + query_size) / x.shape[0]))
            x = np.tile(x, (reps, 1))
        rng = np.random.default_rng(seed)
        idx = rng.permutation(x.shape[0])[: db_size + query_size]
        xb = project_to_binary256(x[idx], seed)
        return xb[:db_size], xb[db_size : db_size + query_size], dataset_id

    if source == "local_npy":
        path = Path(str(dataset_cfg["path"]))
        arr = np.load(path)
        if arr.ndim == 2 and arr.shape[1] == 32 and arr.dtype == np.uint8:
            x = arr
        else:
            x = project_to_binary256(arr, seed)
        if x.shape[0] < db_size + query_size:
            reps = int(math.ceil((db_size + query_size) / x.shape[0]))
            x = np.tile(x, (reps, 1))
        return x[:db_size], x[db_size : db_size + query_size], dataset_id

    raise ValueError(f"Unknown dataset source: {source}")


BITCOUNT_TABLE = np.array([bin(i).count("1") for i in range(256)], dtype=np.uint8)


def hamming_distances(query_bytes: np.ndarray, db_bytes: np.ndarray) -> np.ndarray:
    xor = np.bitwise_xor(db_bytes, query_bytes)
    return BITCOUNT_TABLE[xor].sum(axis=1).astype(np.int32)


def bruteforce_topk(db: np.ndarray, queries: np.ndarray, k: int, max_queries: Optional[int]) -> Tuple[np.ndarray, np.ndarray]:
    qn = queries.shape[0] if max_queries is None else min(queries.shape[0], max_queries)
    ids = np.zeros((qn, k), dtype=np.int64)
    dists = np.zeros((qn, k), dtype=np.int32)
    for i in range(qn):
        dist = hamming_distances(queries[i], db)
        order = np.argpartition(dist, kth=min(k - 1, len(dist) - 1))[:k]
        order = order[np.argsort(dist[order])]
        ids[i, : len(order)] = order
        dists[i, : len(order)] = dist[order]
    return ids, dists


def run_faiss_baseline(db: np.ndarray, queries: np.ndarray, k: int) -> Dict[str, Any]:
    import faiss  # type: ignore

    index = faiss.IndexBinaryFlat(256)
    t0 = dt.datetime.now().timestamp()
    index.add(db)
    build_s = dt.datetime.now().timestamp() - t0

    lat_ms: List[float] = []
    all_ids = np.zeros((queries.shape[0], k), dtype=np.int64)
    all_d = np.zeros((queries.shape[0], k), dtype=np.int64)

    t1 = dt.datetime.now().timestamp()
    for i in range(queries.shape[0]):
        q = queries[i : i + 1]
        q0 = dt.datetime.now().timestamp()
        d, ids = index.search(q, k)
        q1 = dt.datetime.now().timestamp()
        lat_ms.append((q1 - q0) * 1000.0)
        all_ids[i] = ids[0]
        all_d[i] = d[0]
    qtime_s = dt.datetime.now().timestamp() - t1
    qps = queries.shape[0] / qtime_s if qtime_s > 0 else 0.0

    return {
        "algorithm": "faiss-binary-flat",
        "framework": "faiss",
        "build_time_s": build_s,
        "qps": qps,
        "lat_p50_ms": percentile(lat_ms, 0.50),
        "lat_p95_ms": percentile(lat_ms, 0.95),
        "lat_p99_ms": percentile(lat_ms, 0.99),
        "ids": all_ids,
        "distances": all_d,
        "index_bytes": int(db.nbytes),
    }


def run_exact_reference_baseline(db: np.ndarray, queries: np.ndarray, k: int) -> Dict[str, Any]:
    lat_ms: List[float] = []
    all_ids = np.zeros((queries.shape[0], k), dtype=np.int64)
    all_d = np.zeros((queries.shape[0], k), dtype=np.int32)

    t0 = dt.datetime.now().timestamp()
    for i in range(queries.shape[0]):
        q0 = dt.datetime.now().timestamp()
        dist = hamming_distances(queries[i], db)
        order = np.argpartition(dist, kth=min(k - 1, len(dist) - 1))[:k]
        order = order[np.argsort(dist[order])]
        all_ids[i, : len(order)] = order
        all_d[i, : len(order)] = dist[order]
        q1 = dt.datetime.now().timestamp()
        lat_ms.append((q1 - q0) * 1000.0)
    total_s = dt.datetime.now().timestamp() - t0
    qps = queries.shape[0] / total_s if total_s > 0 else 0.0

    return {
        "algorithm": "reference-exact-hamming",
        "framework": "internal_exact",
        "build_time_s": 0.0,
        "qps": qps,
        "lat_p50_ms": percentile(lat_ms, 0.50),
        "lat_p95_ms": percentile(lat_ms, 0.95),
        "lat_p99_ms": percentile(lat_ms, 0.99),
        "ids": all_ids,
        "distances": all_d,
        "index_bytes": int(db.nbytes),
    }


def run_hnswlib_baseline(db: np.ndarray, queries: np.ndarray, k: int, seed: int) -> Dict[str, Any]:
    import hnswlib  # type: ignore

    db_float = np.unpackbits(db, axis=1, bitorder="big").astype(np.float32)
    q_float = np.unpackbits(queries, axis=1, bitorder="big").astype(np.float32)

    index = hnswlib.Index(space="l2", dim=256)
    t0 = dt.datetime.now().timestamp()
    index.init_index(max_elements=db_float.shape[0], ef_construction=200, M=16, random_seed=seed)
    index.add_items(db_float, np.arange(db_float.shape[0]))
    index.set_ef(max(50, k * 4))
    build_s = dt.datetime.now().timestamp() - t0

    lat_ms: List[float] = []
    all_ids = np.zeros((q_float.shape[0], k), dtype=np.int64)
    all_d = np.zeros((q_float.shape[0], k), dtype=np.float32)

    t1 = dt.datetime.now().timestamp()
    for i in range(q_float.shape[0]):
        q0 = dt.datetime.now().timestamp()
        ids, d = index.knn_query(q_float[i : i + 1], k=k)
        q1 = dt.datetime.now().timestamp()
        lat_ms.append((q1 - q0) * 1000.0)
        all_ids[i] = ids[0]
        all_d[i] = d[0]
    qtime_s = dt.datetime.now().timestamp() - t1
    qps = q_float.shape[0] / qtime_s if qtime_s > 0 else 0.0

    return {
        "algorithm": "hnswlib-l2-bits",
        "framework": "hnswlib",
        "build_time_s": build_s,
        "qps": qps,
        "lat_p50_ms": percentile(lat_ms, 0.50),
        "lat_p95_ms": percentile(lat_ms, 0.95),
        "lat_p99_ms": percentile(lat_ms, 0.99),
        "ids": all_ids,
        "distances": all_d,
        "index_bytes": int(db.nbytes),
    }


def run_nmslib_baseline(db: np.ndarray, queries: np.ndarray, k: int) -> Dict[str, Any]:
    import nmslib  # type: ignore

    # `bit_hamming` expects vectors interpreted as strings/bytes depending on binding.
    # To keep runtime stable across versions, we convert to unpacked {0,1} and use l2.
    db_float = np.unpackbits(db, axis=1, bitorder="big").astype(np.float32)
    q_float = np.unpackbits(queries, axis=1, bitorder="big").astype(np.float32)

    index = nmslib.init(method="hnsw", space="l2")
    t0 = dt.datetime.now().timestamp()
    index.addDataPointBatch(db_float)
    index.createIndex({"post": 2}, print_progress=False)
    build_s = dt.datetime.now().timestamp() - t0

    lat_ms: List[float] = []
    all_ids = np.zeros((q_float.shape[0], k), dtype=np.int64)
    all_d = np.zeros((q_float.shape[0], k), dtype=np.float32)
    t1 = dt.datetime.now().timestamp()
    for i in range(q_float.shape[0]):
        q0 = dt.datetime.now().timestamp()
        ids, d = index.knnQuery(q_float[i], k=k)
        q1 = dt.datetime.now().timestamp()
        lat_ms.append((q1 - q0) * 1000.0)
        ids = np.asarray(ids, dtype=np.int64)
        d = np.asarray(d, dtype=np.float32)
        if ids.shape[0] < k:
            pad = np.full((k - ids.shape[0],), -1, dtype=np.int64)
            ids = np.concatenate([ids, pad], axis=0)
            d = np.concatenate([d, np.full((k - d.shape[0],), np.inf, dtype=np.float32)], axis=0)
        all_ids[i] = ids[:k]
        all_d[i] = d[:k]
    qtime_s = dt.datetime.now().timestamp() - t1
    qps = q_float.shape[0] / qtime_s if qtime_s > 0 else 0.0

    return {
        "algorithm": "nmslib-hnsw-l2-bits",
        "framework": "nmslib",
        "build_time_s": build_s,
        "qps": qps,
        "lat_p50_ms": percentile(lat_ms, 0.50),
        "lat_p95_ms": percentile(lat_ms, 0.95),
        "lat_p99_ms": percentile(lat_ms, 0.99),
        "ids": all_ids,
        "distances": all_d,
        "index_bytes": int(db.nbytes),
    }


def evaluate_ann_result(db: np.ndarray,
                        queries: np.ndarray,
                        retrieved_ids: np.ndarray,
                        exact_topk_ids: Optional[np.ndarray],
                        exact_topk_dists: Optional[np.ndarray],
                        audit_every: int,
                        audit_max_queries: int) -> Dict[str, Any]:
    qn = queries.shape[0]
    k = retrieved_ids.shape[1]
    min_hamming_vals: List[float] = []
    recall_id_vals: List[float] = []
    recall_tie_vals: List[float] = []
    gaps: List[float] = []

    for i in range(qn):
        ids_raw = retrieved_ids[i]
        ids = ids_raw[(ids_raw >= 0) & (ids_raw < db.shape[0])]
        if ids.size == 0:
            continue
        cand = db[ids]
        pred_dist = hamming_distances(queries[i], cand)
        min_hamming_vals.append(float(pred_dist.min()))

    if exact_topk_ids is not None and exact_topk_dists is not None:
        eval_q = min(exact_topk_ids.shape[0], qn)
        for i in range(eval_q):
            pred_ids = [int(x) for x in retrieved_ids[i] if 0 <= int(x) < db.shape[0]]
            gt = set(int(x) for x in exact_topk_ids[i] if x >= 0)
            pred = set(pred_ids)
            if gt:
                recall_id_vals.append(len(gt & pred) / max(1, k))

            if pred_ids:
                pred_arr = np.array(pred_ids, dtype=np.int64)
                pred_min_all = hamming_distances(queries[i], db[pred_arr])
                if exact_topk_dists.shape[1] > 0:
                    gt_kth_idx = min(k - 1, exact_topk_dists.shape[1] - 1)
                    gt_kth = float(exact_topk_dists[i][gt_kth_idx])
                    tie_hits = float(np.sum(pred_min_all <= gt_kth))
                    recall_tie_vals.append(tie_hits / max(1, k))
                pred_min = float(pred_min_all.min())
            if i % max(1, audit_every) == 0 and len(gaps) < audit_max_queries:
                if pred_ids:
                    pred_min = float(pred_min)
                else:
                    pred_min = 256.0
                gt_min = float(exact_topk_dists[i][0])
                gaps.append(abs(pred_min - gt_min))

    recall_id = numeric_mean(recall_id_vals)
    recall_tie = numeric_mean(recall_tie_vals)
    recall_final = recall_tie if recall_tie is not None else recall_id
    return {
        "recall_at_k": recall_final,
        "recall_at_k_id_overlap": recall_id,
        "recall_at_k_tie_aware": recall_tie,
        "error_bits_min_hamming": numeric_mean(min_hamming_vals),
        "hamming_audit_gap": numeric_mean(gaps),
    }


def run_ann_track(repo_root: Path,
                  cfg: Dict[str, Any],
                  run_id: str,
                  tier: str,
                  scale: str,
                  seed: int,
                  scale_profile: Dict[str, Any]) -> List[Dict[str, Any]]:
    rows: List[Dict[str, Any]] = []
    ann_cfg = cfg["tracks"]["ann"]
    if not ann_cfg.get("enabled", True):
        return rows

    db_size = int(scale_profile["ann_db"])
    query_size = int(scale_profile["ann_queries"])
    top_k = int(cfg["global"]["top_k"])
    max_gt_q = int(ann_cfg.get("max_ground_truth_queries", cfg["global"]["hamming_audit_max_queries"]))

    for dsi, ds_cfg in enumerate(ann_cfg.get("datasets", [])):
        ds_seed = (seed + dsi * 9973) & 0xFFFFFFFFFFFFFFFF
        dataset_id = str(ds_cfg.get("id", f"dataset_{dsi}"))
        try:
            db, queries, source_id = load_dataset(ds_cfg, db_size, query_size, ds_seed)
        except Exception as exc:
            for baseline in ann_cfg.get("baselines", []):
                row = base_row(run_id, seed, "ann", tier, scale, dataset_id, f"{baseline}/unavailable")
                row["framework"] = ann_cfg.get("framework")
                row["status"] = "failed"
                row["reason"] = f"dataset load error: {exc}"
                row["source_id"] = dataset_id
                row["requested_scale_points"] = int(scale_profile["requested_points"])
                row["effective_db_size"] = db_size
                row["effective_query_size"] = query_size
                rows.append(ensure_result_shape(row))
            continue

        exact_ids, exact_d = bruteforce_topk(db, queries, top_k, max_gt_q)

        ref_queries = queries[: exact_ids.shape[0]]
        ref_payload = run_exact_reference_baseline(db, ref_queries, top_k)
        ref_eval = evaluate_ann_result(
            db,
            ref_queries,
            np.asarray(ref_payload["ids"]),
            exact_ids,
            exact_d,
            int(cfg["global"]["hamming_audit_every"]),
            int(cfg["global"]["hamming_audit_max_queries"]),
        )
        ref_row = base_row(run_id, seed, "ann", tier, scale, dataset_id, str(ref_payload["algorithm"]))
        ref_row["framework"] = str(ref_payload["framework"])
        ref_row["source_id"] = source_id
        ref_row["requested_scale_points"] = int(scale_profile["requested_points"])
        ref_row["effective_db_size"] = int(db.shape[0])
        ref_row["effective_query_size"] = int(ref_queries.shape[0])
        ref_row["build_time_s"] = float(ref_payload["build_time_s"])
        ref_row["qps"] = float(ref_payload["qps"])
        ref_row["lat_p50_ms"] = float(ref_payload["lat_p50_ms"])
        ref_row["lat_p95_ms"] = float(ref_payload["lat_p95_ms"])
        ref_row["lat_p99_ms"] = float(ref_payload["lat_p99_ms"])
        ref_row["rss_mb"] = rss_mb()
        ref_row["index_bytes"] = int(ref_payload["index_bytes"])
        ref_row["bytes_per_insert"] = float(ref_row["index_bytes"]) / max(1, int(db.shape[0]))
        ref_row["recall_at_k"] = ref_eval["recall_at_k"]
        ref_row["recall_at_k_id_overlap"] = ref_eval["recall_at_k_id_overlap"]
        ref_row["recall_at_k_tie_aware"] = ref_eval["recall_at_k_tie_aware"]
        ref_row["update_latency_ms"] = None
        ref_row["energy_proxy"] = None
        ref_row["error_bits_min_hamming"] = ref_eval["error_bits_min_hamming"]
        ref_row["error_bits_prefix_gap"] = None
        ref_row["hamming_audit_gap"] = ref_eval["hamming_audit_gap"]
        rows.append(ensure_result_shape(ref_row))

        baselines = ann_cfg.get("baselines", [])
        for baseline in baselines:
            row = base_row(run_id, seed, "ann", tier, scale, dataset_id, baseline)
            row["framework"] = ann_cfg.get("framework")
            row["source_id"] = source_id
            row["requested_scale_points"] = int(scale_profile["requested_points"])
            row["effective_db_size"] = int(db.shape[0])
            row["effective_query_size"] = int(queries.shape[0])
            row["rss_mb"] = rss_mb()
            result_payload: Optional[Dict[str, Any]] = None

            try:
                if baseline == "faiss":
                    result_payload = run_faiss_baseline(db, queries, top_k)
                elif baseline == "hnswlib":
                    result_payload = run_hnswlib_baseline(db, queries, top_k, ds_seed)
                elif baseline == "nmslib":
                    result_payload = run_nmslib_baseline(db, queries, top_k)
                else:
                    row["status"] = "skipped"
                    row["reason"] = f"unknown baseline '{baseline}'"
            except Exception as exc:
                row["status"] = "skipped"
                row["reason"] = f"{baseline} unavailable or failed: {exc}"

            if result_payload is None:
                rows.append(ensure_result_shape(row))
                continue

            eval_payload = evaluate_ann_result(
                db,
                queries,
                np.asarray(result_payload["ids"]),
                exact_ids,
                exact_d,
                int(cfg["global"]["hamming_audit_every"]),
                int(cfg["global"]["hamming_audit_max_queries"]),
            )

            row["algorithm"] = str(result_payload["algorithm"])
            row["framework"] = str(result_payload["framework"])
            row["build_time_s"] = float(result_payload["build_time_s"])
            row["qps"] = float(result_payload["qps"])
            row["lat_p50_ms"] = float(result_payload["lat_p50_ms"])
            row["lat_p95_ms"] = float(result_payload["lat_p95_ms"])
            row["lat_p99_ms"] = float(result_payload["lat_p99_ms"])
            row["index_bytes"] = int(result_payload["index_bytes"])
            row["bytes_per_insert"] = float(row["index_bytes"]) / max(1, int(db.shape[0]))
            row["recall_at_k"] = eval_payload["recall_at_k"]
            row["recall_at_k_id_overlap"] = eval_payload["recall_at_k_id_overlap"]
            row["recall_at_k_tie_aware"] = eval_payload["recall_at_k_tie_aware"]
            row["update_latency_ms"] = None
            row["energy_proxy"] = None
            row["error_bits_min_hamming"] = eval_payload["error_bits_min_hamming"]
            row["error_bits_prefix_gap"] = None
            row["hamming_audit_gap"] = eval_payload["hamming_audit_gap"]
            rows.append(ensure_result_shape(row))

    return rows


def run_online_track(repo_root: Path,
                     cfg: Dict[str, Any],
                     run_id: str,
                     tier: str,
                     scale: str,
                     seed: int,
                     scale_profile: Dict[str, Any]) -> List[Dict[str, Any]]:
    rows: List[Dict[str, Any]] = []
    online_cfg = cfg["tracks"]["online"]
    if not online_cfg.get("enabled", True):
        return rows

    tool = repo_root / cfg["core_tool_bin"]
    events = int(scale_profile["online_events"])
    delay = int(online_cfg.get("delay", 64))
    drift_interval = int(online_cfg.get("drift_interval", 5000))
    sample_rate = int(online_cfg.get("hamming_sample_rate", 10))
    buffer_cap = int(online_cfg.get("buffer_cap", 65536))

    cmd = [
        str(tool),
        "online",
        "--events",
        str(events),
        "--delay",
        str(delay),
        "--drift-interval",
        str(drift_interval),
        "--seed",
        str(seed),
        "--hamming-sample-rate",
        str(sample_rate),
        "--buffer-cap",
        str(buffer_cap),
    ]
    cp = safe_run(cmd, cwd=repo_root)
    if cp.returncode != 0:
        row = base_row(run_id, seed, "online", tier, scale, "online:synthetic_drift", "core256/online_delayed")
        row["framework"] = online_cfg.get("framework")
        row["status"] = "failed"
        row["reason"] = cp.stderr.strip() or cp.stdout.strip() or f"exit={cp.returncode}"
        row["requested_scale_points"] = int(scale_profile["requested_points"])
        row["effective_db_size"] = None
        row["effective_query_size"] = events
        rows.append(ensure_result_shape(row))
        return rows

    payload = None
    for line in cp.stdout.splitlines():
        s = line.strip()
        if s.startswith("{"):
            payload = json.loads(s)
            break
    if payload is None:
        row = base_row(run_id, seed, "online", tier, scale, "online:synthetic_drift", "core256/online_delayed")
        row["framework"] = online_cfg.get("framework")
        row["status"] = "failed"
        row["reason"] = "no JSON payload from online tool"
        rows.append(ensure_result_shape(row))
        return rows

    row = base_row(run_id, seed, "online", tier, scale, "online:synthetic_drift", "core256/online_delayed")
    row["framework"] = online_cfg.get("framework")
    row["operation"] = "classifier:prefix_match"
    row["requested_scale_points"] = int(scale_profile["requested_points"])
    row["effective_query_size"] = int(payload.get("events", events))
    row["build_time_s"] = float(payload.get("total_time_s", 0.0))
    row["qps"] = float(payload.get("qps", 0.0))
    row["lat_p50_ms"] = float(payload.get("update_lat_p50_ms", 0.0))
    row["lat_p95_ms"] = float(payload.get("update_lat_p95_ms", 0.0))
    row["lat_p99_ms"] = float(payload.get("update_lat_p99_ms", 0.0))
    row["rss_mb"] = rss_mb()
    row["index_bytes"] = int(payload.get("index_bytes", 0))
    row["bytes_per_insert"] = float(payload.get("bytes_per_insert", 0.0))
    row["recall_at_k"] = None
    row["update_latency_ms"] = float(payload.get("update_latency_ms", 0.0))
    row["energy_proxy"] = None
    row["error_bits_min_hamming"] = float(payload.get("error_bits_min_hamming", -1.0))
    row["error_bits_prefix_gap"] = float(payload.get("error_bits_prefix_gap", -1.0))
    row["hamming_audit_gap"] = float(payload.get("hamming_audit_gap", -1.0))
    row["online_accuracy"] = float(payload.get("accuracy", 0.0))
    row["online_f1"] = float(payload.get("f1", 0.0))
    row["adaptation_lag_steps"] = float(payload.get("adaptation_lag_steps", -1.0))
    classifier_mode = str(payload.get("classifier_mode", "prefix_match"))
    row["operation"] = f"classifier:{classifier_mode}"
    rows.append(ensure_result_shape(row))
    return rows


def run_tracks(repo_root: Path,
               cfg: Dict[str, Any],
               run_id: str,
               tracks: Sequence[str],
               tiers: Sequence[str],
               scales: Sequence[str],
               seed: int) -> List[Dict[str, Any]]:
    rows: List[Dict[str, Any]] = []
    for tier in tiers:
        for scale in scales:
            profile = cfg["scale_profiles"][scale]
            for track in tracks:
                if track == "core":
                    rows.extend(run_core_track(repo_root, cfg, run_id, tier, scale, seed, profile))
                elif track == "ann":
                    rows.extend(run_ann_track(repo_root, cfg, run_id, tier, scale, seed, profile))
                elif track == "online":
                    rows.extend(run_online_track(repo_root, cfg, run_id, tier, scale, seed, profile))
                else:
                    raise ValueError(f"unknown track: {track}")
    return rows


def make_manifest(repo_root: Path, cfg: Dict[str, Any], run_id: str, config_path: Path) -> Dict[str, Any]:
    hw = hardware_fingerprint()
    return {
        "run_id": run_id,
        "created_utc": utc_now(),
        "config_path": str(config_path),
        "config_hash": hash_jsonable(cfg),
        "git_commit": detect_git_commit(repo_root),
        "container_hash": detect_container_hash(),
        "hardware_fingerprint": hw,
        "hardware_fingerprint_hash": hash_jsonable(hw),
        "tool_versions": {
            "python": sys.version.splitlines()[0],
            "gcc": parse_tool_version(["gcc", "--version"]),
            "bench_core256_tool": parse_tool_version([str(repo_root / cfg["core_tool_bin"]), "--help"]),
        },
        "official_source_pack": OFFICIAL_SOURCE_PACK,
        "licensing_policy": cfg.get("licensing_policy", "internal_review"),
        "legal_gate_status": "pending_internal_review",
    }


def write_legal_checklist(run_dir: Path, cfg: Dict[str, Any]) -> None:
    path = run_dir / "legal_review_checklist.csv"
    with path.open("w", encoding="utf-8", newline="") as f:
        w = csv.DictWriter(f, fieldnames=["source", "type", "status", "notes"])
        w.writeheader()
        for src in OFFICIAL_SOURCE_PACK:
            w.writerow({"source": src, "type": "official_source", "status": "pending", "notes": cfg.get("licensing_policy")})
        for ds in cfg.get("tracks", {}).get("ann", {}).get("datasets", []):
            w.writerow({"source": ds.get("id"), "type": ds.get("source"), "status": "pending", "notes": "dataset legal review required"})


def write_evidence_matrix_csv(path: Path, rows: Sequence[Dict[str, Any]]) -> None:
    fields = [
        "run_id",
        "track",
        "dataset_id",
        "tier",
        "scale",
        "algorithm",
        "status",
        "framework",
        "qps",
        "lat_p50_ms",
        "lat_p95_ms",
        "lat_p99_ms",
        "build_time_s",
        "rss_mb",
        "index_bytes",
        "bytes_per_insert",
        "recall_at_k",
        "recall_at_k_id_overlap",
        "recall_at_k_tie_aware",
        "error_bits_min_hamming",
        "error_bits_prefix_gap",
        "hamming_audit_gap",
        "online_accuracy",
        "online_f1",
        "update_latency_ms",
        "adaptation_lag_steps",
        "reason",
    ]
    sorted_rows = sorted(
        rows,
        key=lambda r: (
            str(r.get("track")),
            str(r.get("dataset_id")),
            str(r.get("tier")),
            str(r.get("scale")),
            str(r.get("algorithm")),
            str(r.get("status")),
        ),
    )
    with path.open("w", encoding="utf-8", newline="") as f:
        w = csv.DictWriter(f, fieldnames=fields)
        w.writeheader()
        for row in sorted_rows:
            out = {k: row.get(k) for k in fields}
            w.writerow(out)


def _pareto_front_flags(entries: Sequence[Dict[str, Any]],
                        maximize_fields: Sequence[str],
                        minimize_fields: Sequence[str]) -> List[bool]:
    flags = [True] * len(entries)
    for i, a in enumerate(entries):
        for j, b in enumerate(entries):
            if i == j:
                continue
            better_or_equal = True
            strictly_better = False
            for field in maximize_fields:
                va = to_float(a.get(field))
                vb = to_float(b.get(field))
                if va is None or vb is None:
                    better_or_equal = False
                    break
                if vb < va:
                    better_or_equal = False
                    break
                if vb > va:
                    strictly_better = True
            if not better_or_equal:
                continue
            for field in minimize_fields:
                va = to_float(a.get(field))
                vb = to_float(b.get(field))
                if va is None or vb is None:
                    better_or_equal = False
                    break
                if vb > va:
                    better_or_equal = False
                    break
                if vb < va:
                    strictly_better = True
            if better_or_equal and strictly_better:
                flags[i] = False
                break
    return flags


def write_ann_pareto_csv(path: Path,
                         rows: Sequence[Dict[str, Any]],
                         maximize_fields: Sequence[str],
                         minimize_fields: Sequence[str],
                         value_fields: Sequence[str]) -> None:
    ann_rows = [
        r for r in rows
        if r.get("track") == "ann"
        and r.get("status") == "ok"
        and to_float(r.get("recall_at_k")) is not None
    ]
    grouped: Dict[Tuple[str, str, str], List[Dict[str, Any]]] = {}
    for row in ann_rows:
        key = (str(row.get("dataset_id")), str(row.get("tier")), str(row.get("scale")))
        grouped.setdefault(key, []).append(row)

    fields = ["dataset_id", "tier", "scale", "algorithm", "status", "is_pareto"] + list(value_fields)
    with path.open("w", encoding="utf-8", newline="") as f:
        w = csv.DictWriter(f, fieldnames=fields)
        w.writeheader()
        for key in sorted(grouped.keys()):
            group_rows = grouped[key]
            flags = _pareto_front_flags(group_rows, maximize_fields=maximize_fields, minimize_fields=minimize_fields)
            for row, flag in zip(group_rows, flags):
                out = {
                    "dataset_id": key[0],
                    "tier": key[1],
                    "scale": key[2],
                    "algorithm": row.get("algorithm"),
                    "status": row.get("status"),
                    "is_pareto": "yes" if flag else "no",
                }
                for vf in value_fields:
                    out[vf] = row.get(vf)
                w.writerow(out)


def compute_acceptance_gates(rows: Sequence[Dict[str, Any]]) -> Dict[str, Any]:
    ok_rows = [r for r in rows if r.get("status") == "ok"]
    now = utc_now()

    latency_viol = 0
    negative_resource = 0
    for r in ok_rows:
        p50 = to_float(r.get("lat_p50_ms"))
        p95 = to_float(r.get("lat_p95_ms"))
        p99 = to_float(r.get("lat_p99_ms"))
        if p50 is not None and p95 is not None and p99 is not None and not (p50 <= p95 <= p99):
            latency_viol += 1
        for f in ("rss_mb", "index_bytes", "bytes_per_insert"):
            v = to_float(r.get(f))
            if v is not None and v < 0:
                negative_resource += 1

    audit_bound = 2.0
    audited_rows = [
        r for r in ok_rows
        if r.get("track") == "ann" and str(r.get("algorithm", "")).startswith("reference-exact-hamming")
    ]
    audit_vals = [to_float(r.get("hamming_audit_gap")) for r in audited_rows]
    audit_viol = sum(1 for v in audit_vals if v is not None and v > audit_bound)

    correctness_pass = latency_viol == 0 and negative_resource == 0 and audit_viol == 0

    repro_groups: Dict[Tuple[str, str, str, str, str], List[Dict[str, Any]]] = {}
    for r in ok_rows:
        key = (
            str(r.get("track")),
            str(r.get("dataset_id")),
            str(r.get("tier")),
            str(r.get("scale")),
            str(r.get("algorithm")),
        )
        repro_groups.setdefault(key, []).append(r)
    repro_group_cvs: List[float] = []
    for group_rows in repro_groups.values():
        if len(group_rows) < 2:
            continue
        qps_cv = coefficient_of_variation([to_float(r.get("qps")) for r in group_rows])
        p95_cv = coefficient_of_variation([to_float(r.get("lat_p95_ms")) for r in group_rows])
        if qps_cv is not None:
            repro_group_cvs.append(qps_cv)
        if p95_cv is not None:
            repro_group_cvs.append(p95_cv)
    repro_cv_max = max(repro_group_cvs) if repro_group_cvs else None
    repro_threshold = 0.25
    if repro_cv_max is None:
        repro_status = "insufficient_data"
    else:
        repro_status = "pass" if repro_cv_max <= repro_threshold else "fail"

    required_tiers = set(TIER_CHOICES)
    required_tracks = set(TRACK_CHOICES)
    tier_track_map: Dict[str, set[str]] = {t: set() for t in required_tiers}
    for r in ok_rows:
        tier = str(r.get("tier"))
        track = str(r.get("track"))
        if tier in tier_track_map and track in required_tracks:
            tier_track_map[tier].add(track)
    tiers_with_all_tracks = sorted([t for t, tracks in tier_track_map.items() if tracks == required_tracks])
    investor_pass = len(tiers_with_all_tracks) == len(required_tiers)

    ok_scales = sorted({str(r.get("scale")) for r in ok_rows if r.get("scale") is not None})
    has_1m = "1M" in ok_scales
    has_10m = "10M" in ok_scales
    has_100m = "100M" in ok_scales
    scale_status = "pass" if has_1m and has_10m else "fail"

    failing = []
    if not correctness_pass:
        failing.append("correctness")
    if repro_status == "fail":
        failing.append("reproducibility")
    if not investor_pass:
        failing.append("investor")
    if scale_status == "fail":
        failing.append("scale")

    overall = "pass" if not failing else "fail"
    return {
        "generated_utc": now,
        "overall_status": overall,
        "failing_gates": failing,
        "gates": {
            "correctness": {
                "status": "pass" if correctness_pass else "fail",
                "latency_monotonic_violations": latency_viol,
                "negative_resource_violations": negative_resource,
                "audit_reference_rows": len(audited_rows),
                "audit_gap_bound": audit_bound,
                "audit_gap_violations": audit_viol,
            },
            "reproducibility": {
                "status": repro_status,
                "cv_threshold": repro_threshold,
                "max_observed_cv": repro_cv_max,
                "groups_with_repetition": sum(1 for g in repro_groups.values() if len(g) >= 2),
            },
            "investor": {
                "status": "pass" if investor_pass else "fail",
                "required_tiers": sorted(required_tiers),
                "required_tracks": sorted(required_tracks),
                "tiers_with_all_tracks": tiers_with_all_tracks,
                "tier_track_coverage": {k: sorted(v) for k, v in tier_track_map.items()},
            },
            "scale": {
                "status": scale_status,
                "ok_scales": ok_scales,
                "required_minimum": ["1M", "10M"],
                "required_final_dossier": "100M",
                "final_dossier_ready": has_100m,
            },
        },
    }


def write_csv_rows(path: Path, rows: Sequence[Dict[str, Any]], fieldnames: Optional[Sequence[str]] = None) -> None:
    if not rows:
        with path.open("w", encoding="utf-8", newline="") as f:
            w = csv.DictWriter(f, fieldnames=list(fieldnames) if fieldnames else ["empty"])
            w.writeheader()
        return
    if fieldnames is None:
        fieldnames = list(rows[0].keys())
    with path.open("w", encoding="utf-8", newline="") as f:
        w = csv.DictWriter(f, fieldnames=list(fieldnames))
        w.writeheader()
        for row in rows:
            w.writerow({k: row.get(k) for k in fieldnames})


def build_memory_speed_rows(rows: Sequence[Dict[str, Any]]) -> List[Dict[str, Any]]:
    out: List[Dict[str, Any]] = []
    for r in rows:
        if r.get("status") != "ok":
            continue
        qps = to_float(r.get("qps"))
        p95 = to_float(r.get("lat_p95_ms"))
        if qps is None and p95 is None:
            continue
        idx_b = to_float(r.get("index_bytes"))
        idx_mb = (idx_b / (1024.0 * 1024.0)) if idx_b is not None else None
        rss = to_float(r.get("rss_mb"))
        row = {
            "track": r.get("track"),
            "dataset_id": r.get("dataset_id"),
            "tier": r.get("tier"),
            "scale": r.get("scale"),
            "algorithm": r.get("algorithm"),
            "status": r.get("status"),
            "qps": qps,
            "lat_p95_ms": p95,
            "index_mb": idx_mb,
            "rss_mb": rss,
            "bytes_per_insert": to_float(r.get("bytes_per_insert")),
            "recall_at_k": to_float(r.get("recall_at_k")),
            "error_bits_min_hamming": to_float(r.get("error_bits_min_hamming")),
            "qps_per_index_mb": (qps / idx_mb) if (qps is not None and idx_mb is not None and idx_mb > 0.0) else None,
            "qps_per_rss_mb": (qps / rss) if (qps is not None and rss is not None and rss > 0.0) else None,
        }
        out.append(row)
    out.sort(
        key=lambda x: (
            str(x.get("track")),
            str(x.get("dataset_id")),
            str(x.get("tier")),
            str(x.get("scale")),
            -(x.get("qps") if isinstance(x.get("qps"), (int, float)) else -1.0),
        )
    )
    return out


def classify_data_regime(dataset_id: str) -> str:
    if dataset_id.startswith("core:"):
        suffix = dataset_id.split(":", 1)[1]
        if suffix == "random":
            return "uniform-random 256-bit"
        if suffix == "clustered":
            return "near-duplicate clustered keys"
        if suffix == "adversarial":
            return "prefix-collision adversarial"
        if suffix == "replay":
            return "replay/hash-stream"
        return f"core-{suffix}"
    if dataset_id == "synthetic_binary256":
        return "binary clustered synthetic"
    if dataset_id.startswith("openml:"):
        return "mixed tabular (numeric + categorical)"
    if dataset_id.startswith("online:"):
        return "stream with concept drift"
    return "custom/unknown"


def build_data_impact_rows(rows: Sequence[Dict[str, Any]]) -> List[Dict[str, Any]]:
    groups: Dict[Tuple[str, str, str, str], List[Dict[str, Any]]] = {}
    for r in rows:
        if r.get("status") != "ok":
            continue
        key = (
            str(r.get("track")),
            str(r.get("dataset_id")),
            str(r.get("tier")),
            str(r.get("scale")),
        )
        groups.setdefault(key, []).append(r)

    out: List[Dict[str, Any]] = []
    for key in sorted(groups.keys()):
        grp = groups[key]
        qps_vals = [to_float(r.get("qps")) for r in grp]
        p95_vals = [to_float(r.get("lat_p95_ms")) for r in grp]
        bpi_vals = [to_float(r.get("bytes_per_insert")) for r in grp]
        rec_vals = [to_float(r.get("recall_at_k")) for r in grp]
        hamm_vals = [to_float(r.get("error_bits_min_hamming")) for r in grp]
        top_qps_row = max(grp, key=lambda r: to_float(r.get("qps")) or -1.0)
        top_mem_row = max(grp, key=lambda r: ((to_float(r.get("qps")) or 0.0) / max(1e-9, (to_float(r.get("index_bytes")) or float("inf")))))
        out.append(
            {
                "track": key[0],
                "dataset_id": key[1],
                "data_regime": classify_data_regime(key[1]),
                "tier": key[2],
                "scale": key[3],
                "rows_in_group": len(grp),
                "mean_qps": numeric_mean(qps_vals),
                "mean_lat_p95_ms": numeric_mean(p95_vals),
                "mean_bytes_per_insert": numeric_mean(bpi_vals),
                "mean_recall_at_k": numeric_mean(rec_vals),
                "mean_error_bits_min_hamming": numeric_mean(hamm_vals),
                "best_algorithm_by_qps": top_qps_row.get("algorithm"),
                "best_qps": to_float(top_qps_row.get("qps")),
                "best_algorithm_by_compactness": top_mem_row.get("algorithm"),
                "best_qps_per_index_mb": (
                    (to_float(top_mem_row.get("qps")) or 0.0)
                    / max(1e-9, (to_float(top_mem_row.get("index_bytes")) or float("inf")) / (1024.0 * 1024.0))
                    if to_float(top_mem_row.get("index_bytes")) is not None
                    else None
                ),
            }
        )
    return out


def render_memory_speed_report(rows: Sequence[Dict[str, Any]], lang: str) -> str:
    ms_rows = build_memory_speed_rows(rows)
    top_qps = sorted(
        [r for r in ms_rows if to_float(r.get("qps")) is not None],
        key=lambda x: to_float(x.get("qps")) or -1.0,
        reverse=True,
    )[:10]
    top_eff = sorted(
        [r for r in ms_rows if to_float(r.get("qps_per_index_mb")) is not None],
        key=lambda x: to_float(x.get("qps_per_index_mb")) or -1.0,
        reverse=True,
    )[:10]

    lines: List[str] = []
    if lang in ("ru", "both"):
        lines.extend([
            "# Отчет по памяти и скорости (RU)",
            "",
            "## Ключевые наблюдения",
            f"- Всего валидных точек: `{len(ms_rows)}`.",
        ])
        if top_qps:
            best = top_qps[0]
            lines.append(
                f"- Пиковая скорость: `{best.get('algorithm')}` на `{best.get('dataset_id')}` "
                f"(`{best.get('tier')}/{best.get('scale')}`) с `qps={to_float(best.get('qps')):.3f}`."
            )
        if top_eff:
            best = top_eff[0]
            lines.append(
                f"- Лучшая эффективность qps/MB (индекс): `{best.get('algorithm')}` на `{best.get('dataset_id')}` "
                f"с `qps_per_index_mb={to_float(best.get('qps_per_index_mb')):.3f}`."
            )
        lines.append("")
    if lang in ("en", "both"):
        lines.extend([
            "# Memory/Speed Report (EN)",
            "",
            "## Key observations",
            f"- Total valid points: `{len(ms_rows)}`.",
        ])
        if top_qps:
            best = top_qps[0]
            lines.append(
                f"- Peak throughput: `{best.get('algorithm')}` on `{best.get('dataset_id')}` "
                f"(`{best.get('tier')}/{best.get('scale')}`) with `qps={to_float(best.get('qps')):.3f}`."
            )
        if top_eff:
            best = top_eff[0]
            lines.append(
                f"- Best qps/MB efficiency (index): `{best.get('algorithm')}` on `{best.get('dataset_id')}` "
                f"with `qps_per_index_mb={to_float(best.get('qps_per_index_mb')):.3f}`."
            )
        lines.append("")

    lines.append("## Top Throughput")
    lines.append("")
    lines.append("| track | dataset | tier | scale | algorithm | qps | p95 ms | index MB | qps/indexMB |")
    lines.append("|---|---|---|---|---|---:|---:|---:|---:|")
    for r in top_qps:
        lines.append(
            f"| {r.get('track')} | {r.get('dataset_id')} | {r.get('tier')} | {r.get('scale')} | "
            f"{r.get('algorithm')} | {to_float(r.get('qps')):.3f} | {to_float(r.get('lat_p95_ms') or 0.0):.6f} | "
            f"{(to_float(r.get('index_mb')) or 0.0):.3f} | {(to_float(r.get('qps_per_index_mb')) or 0.0):.3f} |"
        )
    lines.append("")
    return "\n".join(lines).strip() + "\n"


def render_data_impact_report(rows: Sequence[Dict[str, Any]], lang: str) -> str:
    impact = build_data_impact_rows(rows)
    core_insert = [r for r in rows if r.get("status") == "ok" and r.get("track") == "core" and r.get("algorithm") == "core256/insert"]
    best_core_speed = max(core_insert, key=lambda r: to_float(r.get("qps")) or -1.0) if core_insert else None
    best_core_compact = min(core_insert, key=lambda r: to_float(r.get("bytes_per_insert")) if to_float(r.get("bytes_per_insert")) is not None else 1e30) if core_insert else None
    ann_rows = [r for r in rows if r.get("status") == "ok" and r.get("track") == "ann" and to_float(r.get("recall_at_k")) is not None]
    ann_99 = [
        r for r in ann_rows
        if (to_float(r.get("recall_at_k")) or 0.0) >= 0.99 and to_float(r.get("lat_p95_ms")) is not None
    ]
    best_ann_recall_latency = min(ann_99, key=lambda r: to_float(r.get("lat_p95_ms")) or 1e30) if ann_99 else None

    lines: List[str] = []
    if lang in ("ru", "both"):
        lines.extend([
            "# Влияние данных на систему (RU)",
            "",
            "## Что лучше работает по данным",
        ])
        if best_core_speed:
            lines.append(
                f"- Для `core256/insert` максимальная скорость на режиме `{best_core_speed.get('dataset_id')}`: "
                f"`qps={(to_float(best_core_speed.get('qps')) or 0.0):.3f}`."
            )
        if best_core_compact:
            lines.append(
                f"- Наиболее компактный режим вставки: `{best_core_compact.get('dataset_id')}` с "
                f"`bytes_per_insert={(to_float(best_core_compact.get('bytes_per_insert')) or 0.0):.3f}`."
            )
        if best_ann_recall_latency:
            lines.append(
                f"- Для ANN при `recall>=0.99` лучший `p95` у `{best_ann_recall_latency.get('algorithm')}` "
                f"на `{best_ann_recall_latency.get('dataset_id')}`: `{(to_float(best_ann_recall_latency.get('lat_p95_ms')) or 0.0):.6f} ms`."
            )
        lines.append("")
    if lang in ("en", "both"):
        lines.extend([
            "# Data Impact Report (EN)",
            "",
            "## What data regimes fit best",
        ])
        if best_core_speed:
            lines.append(
                f"- For `core256/insert`, peak speed is on `{best_core_speed.get('dataset_id')}`: "
                f"`qps={(to_float(best_core_speed.get('qps')) or 0.0):.3f}`."
            )
        if best_core_compact:
            lines.append(
                f"- Most compact insert regime: `{best_core_compact.get('dataset_id')}` with "
                f"`bytes_per_insert={(to_float(best_core_compact.get('bytes_per_insert')) or 0.0):.3f}`."
            )
        if best_ann_recall_latency:
            lines.append(
                f"- For ANN at `recall>=0.99`, best `p95` is `{best_ann_recall_latency.get('algorithm')}` "
                f"on `{best_ann_recall_latency.get('dataset_id')}`: `{(to_float(best_ann_recall_latency.get('lat_p95_ms')) or 0.0):.6f} ms`."
            )
        lines.append("")

    lines.append("## Aggregated Data Regime Table")
    lines.append("")
    lines.append("| track | dataset | regime | tier | scale | mean qps | mean p95 ms | mean bytes/insert | mean recall |")
    lines.append("|---|---|---|---|---|---:|---:|---:|---:|")
    for r in impact[:60]:
        lines.append(
            f"| {r.get('track')} | {r.get('dataset_id')} | {r.get('data_regime')} | {r.get('tier')} | {r.get('scale')} | "
            f"{(to_float(r.get('mean_qps')) or 0.0):.3f} | {(to_float(r.get('mean_lat_p95_ms')) or 0.0):.6f} | "
            f"{(to_float(r.get('mean_bytes_per_insert')) or 0.0):.3f} | {(to_float(r.get('mean_recall_at_k')) or 0.0):.3f} |"
        )
    lines.append("")
    return "\n".join(lines).strip() + "\n"


def render_storage_characteristics_report(rows: Sequence[Dict[str, Any]], lang: str) -> str:
    insert_rows = [
        r for r in rows
        if r.get("status") == "ok" and r.get("track") == "core" and r.get("algorithm") == "core256/insert"
    ]
    bpi_vals = [to_float(r.get("bytes_per_insert")) for r in insert_rows if to_float(r.get("bytes_per_insert")) is not None]
    qps_vals = [to_float(r.get("qps")) for r in insert_rows if to_float(r.get("qps")) is not None]
    bpi_min = min(bpi_vals) if bpi_vals else None
    bpi_max = max(bpi_vals) if bpi_vals else None
    bpi_mean = numeric_mean(bpi_vals)
    qps_mean = numeric_mean(qps_vals)

    lines: List[str] = []
    if lang in ("ru", "both"):
        lines.extend([
            "# Хранение и структурные характеристики (RU)",
            "",
            "Состояние `core256` хранится как CDG-TCA автомат с двумя битовыми строками `(f, s)` одинаковой длины `n_t`.",
            "При конфликте по первому несовпадению `k` рост состояния равен `d = m-k-1`; сериализация выполняется в формате `CDG1`.",
            "Оценка памяти состояния: примерно `2 * n_t` бит полезной нагрузки + служебный заголовок формата.",
            "",
            "## Наблюдения по benchmark",
            f"- `bytes_per_insert`: min={('n/a' if bpi_min is None else f'{bpi_min:.3f}')}, "
            f"mean={('n/a' if bpi_mean is None else f'{bpi_mean:.3f}')}, "
            f"max={('n/a' if bpi_max is None else f'{bpi_max:.3f}')}.",
            f"- Средняя скорость `core256/insert`: {('n/a' if qps_mean is None else f'{qps_mean:.3f} qps')}.",
            "",
        ])
    if lang in ("en", "both"):
        lines.extend([
            "# Storage and Structural Characteristics (EN)",
            "",
            "`core256` stores state as a CDG-TCA automaton with two equal-length bit strings `(f, s)`.",
            "On first mismatch at bit `k`, growth is `d = m-k-1`; serialization is `CDG1`.",
            "Approximate state payload is `2 * n_t` bits plus format header overhead.",
            "",
            "## Observed in benchmark runs",
            f"- `bytes_per_insert`: min={('n/a' if bpi_min is None else f'{bpi_min:.3f}')}, "
            f"mean={('n/a' if bpi_mean is None else f'{bpi_mean:.3f}')}, "
            f"max={('n/a' if bpi_max is None else f'{bpi_max:.3f}')}.",
            f"- Mean `core256/insert` speed: {('n/a' if qps_mean is None else f'{qps_mean:.3f} qps')}.",
            "",
        ])
    return "\n".join(lines).strip() + "\n"


def write_external_benchmark_library_csv(path: Path) -> None:
    write_csv_rows(
        path,
        [dict(x) for x in EXTERNAL_BENCH_LIBRARY],
        fieldnames=("category", "name", "scope", "url"),
    )


def render_external_benchmarks_report(lang: str) -> str:
    lines: List[str] = []
    if lang in ("ru", "both"):
        lines.extend([
            "# Внешние официальные benchmark-наборы (RU)",
            "",
            "Ниже — официальные наборы для валидации памяти/скорости, сжатия логики (truth-table) и деревьев/индексов.",
            "",
        ])
    if lang in ("en", "both"):
        lines.extend([
            "# External Official Benchmark Sets (EN)",
            "",
            "The list below contains official sources for memory/speed validation, truth-table/logic compression, and tree/index workloads.",
            "",
        ])

    lines.append("| category | benchmark | scope | url |")
    lines.append("|---|---|---|---|")
    for entry in EXTERNAL_BENCH_LIBRARY:
        lines.append(
            f"| {entry['category']} | {entry['name']} | {entry['scope']} | {entry['url']} |"
        )
    lines.append("")
    if lang in ("ru", "both"):
        lines.extend([
            "## Как использовать эти бенчи для выбора данных",
            "- ANN-Benchmarks: смотреть фронт `Recall/QPS` и `Recall/Index size` для выбора режима качества/памяти.",
            "- Big-ANN: для практики использовать треки Filtered/OOD/Sparse/Streaming и фиксировать `QPS@90% recall` на нормализованном железе.",
            "- YCSB: использовать `workloada` и `workloadf` (Zipfian) для оценки чувствительности к «горячим» ключам и RMW-паттерну.",
            "- EPFL + table_ex + IWLS2005: использовать для truth-table/logic compression и проверки масштабируемости на arithmetic/control/MtM.",
            "",
        ])
    if lang in ("en", "both"):
        lines.extend([
            "## How to use these benchmarks for data-fit decisions",
            "- ANN-Benchmarks: use `Recall/QPS` and `Recall/Index size` frontiers for quality-memory tradeoff selection.",
            "- Big-ANN: use Filtered/OOD/Sparse/Streaming tracks and compare `QPS@90% recall` on normalized hardware.",
            "- YCSB: run `workloada` and `workloadf` (Zipfian) to evaluate hot-key and read-modify-write sensitivity.",
            "- EPFL + table_ex + IWLS2005: use for truth-table/logic compression and scalability checks on arithmetic/control/MtM circuits.",
            "",
        ])
    return "\n".join(lines).strip() + "\n"


def render_specification_doc(lang: str, manifest: Dict[str, Any]) -> str:
    run_ids = manifest.get("run_ids") or [manifest.get("run_id")]
    source_pack = manifest.get("official_source_pack", OFFICIAL_SOURCE_PACK)
    if lang == "ru":
        lines = [
            "# Спецификация Benchmark Zone (RU)",
            "",
            "## Идентификация",
            f"- run_ids: `{', '.join(str(x) for x in run_ids if x)}`",
            f"- generated_utc: `{utc_now()}`",
            "",
            "## CLI-контракт",
            "- `bench run --config bench_zone.yaml --track {core,ann,online} --tier {server,desktop,edge} --scale {1M,10M,100M}`",
            "- `bench report --input <run_dir...> --output <dossier_dir> --lang {ru,en,both} --format {md,pdf,csv,json,xlsx,docx,all}`",
            "",
            "## Формулы метрик",
            "- `error_bits_min_hamming = min(popcount(query XOR candidate))`",
            "- `error_bits_prefix_gap = 256 - matched_bits`",
            "- `hamming_audit_gap = |pred_min_hamming - exact_min_hamming|` (по выборочным brute-force окнам)",
            "- `recall_at_k` (tie-aware): доля кандидатов с расстоянием `<=` порогу k-го истинного соседа",
            "- `recall_at_k_id_overlap`: классический overlap по id (диагностический)",
            "",
            "## Схема результатов",
            "- Базовые поля: `dataset_id, tier, scale, algorithm, build_time_s, qps, lat_p50_ms, lat_p95_ms, lat_p99_ms, rss_mb, index_bytes, bytes_per_insert, recall_at_k, update_latency_ms, energy_proxy, seed, run_id`",
            "- Ошибки/сходство: `error_bits_min_hamming, error_bits_prefix_gap, hamming_audit_gap`",
            "- Online KPI: `online_accuracy, online_f1, adaptation_lag_steps`",
            "- Доп. отчеты: `dossier_memory_speed_*`, `dossier_data_impact_*`, `dossier_storage_characteristics_*`, `dossier_external_benchmarks_*`",
            "",
            "## Официальный source pack",
        ]
    else:
        lines = [
            "# Benchmark Zone Specification (EN)",
            "",
            "## Identification",
            f"- run_ids: `{', '.join(str(x) for x in run_ids if x)}`",
            f"- generated_utc: `{utc_now()}`",
            "",
            "## CLI contract",
            "- `bench run --config bench_zone.yaml --track {core,ann,online} --tier {server,desktop,edge} --scale {1M,10M,100M}`",
            "- `bench report --input <run_dir...> --output <dossier_dir> --lang {ru,en,both} --format {md,pdf,csv,json,xlsx,docx,all}`",
            "",
            "## Metric formulas",
            "- `error_bits_min_hamming = min(popcount(query XOR candidate))`",
            "- `error_bits_prefix_gap = 256 - matched_bits`",
            "- `hamming_audit_gap = |pred_min_hamming - exact_min_hamming|` (sampled brute-force windows)",
            "- `recall_at_k` (tie-aware): fraction of retrieved candidates with distance `<=` true k-th neighbor distance",
            "- `recall_at_k_id_overlap`: classic id-overlap recall (diagnostic)",
            "",
            "## Result schema",
            "- Core fields: `dataset_id, tier, scale, algorithm, build_time_s, qps, lat_p50_ms, lat_p95_ms, lat_p99_ms, rss_mb, index_bytes, bytes_per_insert, recall_at_k, update_latency_ms, energy_proxy, seed, run_id`",
            "- Error/similarity: `error_bits_min_hamming, error_bits_prefix_gap, hamming_audit_gap`",
            "- Online KPIs: `online_accuracy, online_f1, adaptation_lag_steps`",
            "- Additional reports: `dossier_memory_speed_*`, `dossier_data_impact_*`, `dossier_storage_characteristics_*`, `dossier_external_benchmarks_*`",
            "",
            "## Official source pack",
        ]
    for src in source_pack:
        lines.append(f"- {src}")
    lines.append("")
    return "\n".join(lines).strip() + "\n"


def render_advantages_doc(rows: Sequence[Dict[str, Any]], lang: str) -> str:
    def f3(v: Optional[float]) -> str:
        return "n/a" if v is None else f"{v:.3f}"

    def f6(v: Optional[float]) -> str:
        return "n/a" if v is None else f"{v:.6f}"

    ok_rows = [r for r in rows if r.get("status") == "ok"]
    ann_rows = [r for r in ok_rows if r.get("track") == "ann" and to_float(r.get("recall_at_k")) is not None]
    by_bucket: Dict[Tuple[str, str, str], List[Dict[str, Any]]] = {}
    for r in ann_rows:
        key = (str(r.get("dataset_id")), str(r.get("tier")), str(r.get("scale")))
        by_bucket.setdefault(key, []).append(r)

    lines: List[str] = []
    if lang == "ru":
        lines.extend([
            "# Преимущества и ограничения (RU)",
            "",
            "Ниже выводы только по измеренным данным текущего пакета.",
            "",
            "## ANN: сравнительные итоги",
        ])
    else:
        lines.extend([
            "# Advantages and Limits (EN)",
            "",
            "Statements below are limited to measured values in this package.",
            "",
            "## ANN comparative outcomes",
        ])

    if not by_bucket:
        lines.append("- ANN-данные отсутствуют." if lang == "ru" else "- No ANN rows were available.")
    else:
        for key in sorted(by_bucket.keys()):
            rows_g = by_bucket[key]
            best_recall = max(rows_g, key=lambda r: to_float(r.get("recall_at_k")) or -1.0)
            best_qps = max(rows_g, key=lambda r: to_float(r.get("qps")) or -1.0)
            best_lat = min(rows_g, key=lambda r: to_float(r.get("lat_p95_ms")) if to_float(r.get("lat_p95_ms")) is not None else 1e30)
            if lang == "ru":
                lines.append(
                    f"- `{key[0]} / {key[1]} / {key[2]}`: лучший `recall@k`={f3(to_float(best_recall.get('recall_at_k')))} ({best_recall.get('algorithm')}), "
                    f"лучший `qps`={f3(to_float(best_qps.get('qps')))} ({best_qps.get('algorithm')}), "
                    f"лучший `p95`={f3(to_float(best_lat.get('lat_p95_ms')))} ms ({best_lat.get('algorithm')})."
                )
            else:
                lines.append(
                    f"- `{key[0]} / {key[1]} / {key[2]}`: best `recall@k`={f3(to_float(best_recall.get('recall_at_k')))} ({best_recall.get('algorithm')}), "
                    f"best `qps`={f3(to_float(best_qps.get('qps')))} ({best_qps.get('algorithm')}), "
                    f"best `p95`={f3(to_float(best_lat.get('lat_p95_ms')))} ms ({best_lat.get('algorithm')})."
                )

    online_rows = [r for r in ok_rows if r.get("track") == "online"]
    if lang == "ru":
        lines.extend(["", "## Online-адаптация"])
    else:
        lines.extend(["", "## Online adaptation"])
    if not online_rows:
        lines.append("- Нет online-строк." if lang == "ru" else "- No online rows available.")
    else:
        acc = numeric_mean([to_float(r.get("online_accuracy")) for r in online_rows])
        f1 = numeric_mean([to_float(r.get("online_f1")) for r in online_rows])
        upd = numeric_mean([to_float(r.get("update_latency_ms")) for r in online_rows])
        if lang == "ru":
            lines.append(f"- Средние: accuracy={f3(acc)}, F1={f3(f1)}, update_latency_ms={f6(upd)}.")
        else:
            lines.append(f"- Means: accuracy={f3(acc)}, F1={f3(f1)}, update_latency_ms={f6(upd)}.")

    lines.append("")
    return "\n".join(lines).strip() + "\n"


def sha256_file(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        while True:
            chunk = f.read(1024 * 1024)
            if not chunk:
                break
            h.update(chunk)
    return h.hexdigest()


def write_evidence_manifest(path: Path, files: Sequence[Path], meta: Dict[str, Any]) -> None:
    payload = {
        "generated_utc": utc_now(),
        "files": [],
        "meta": meta,
    }
    for file_path in files:
        if not file_path.exists():
            continue
        payload["files"].append(
            {
                "name": file_path.name,
                "bytes": file_path.stat().st_size,
                "sha256": sha256_file(file_path),
            }
        )
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def _copy_to_path(src: Path, dst: Path) -> bool:
    if not src.exists():
        return False
    dst.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(src, dst)
    return True


def _zip_directory(root: Path, zip_path: Path) -> None:
    import zipfile

    zip_path.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for p in sorted(root.rglob("*")):
            if p.is_file():
                zf.write(p, arcname=str(p.relative_to(root.parent)))


def _collect_files_recursive(root: Path) -> List[Path]:
    files: List[Path] = []
    if not root.exists():
        return files
    for p in sorted(root.rglob("*")):
        if p.is_file():
            files.append(p)
    return files


def build_official_stack_bundle(output_dir: Path,
                                input_dirs: Sequence[Path],
                                lang: str,
                                manifest: Dict[str, Any],
                                report_json: Dict[str, Any],
                                gates: Dict[str, Any],
                                artifacts: Dict[str, Path]) -> Tuple[Path, Path]:
    stack_root = output_dir / "OFFICIAL_STACK"
    stack_root.mkdir(parents=True, exist_ok=True)

    sec_overview = stack_root / "00_OVERVIEW"
    sec_spec = stack_root / "01_SPECIFICATION"
    sec_methods = stack_root / "02_BENCH_METHODS"
    sec_results = stack_root / "03_RESULTS"
    sec_analysis = stack_root / "04_ANALYSIS"
    sec_repro = stack_root / "05_REPRO_LEGAL"
    sec_presentation = stack_root / "06_PRESENTATION"
    for sec in (sec_overview, sec_spec, sec_methods, sec_results, sec_analysis, sec_repro, sec_presentation):
        sec.mkdir(parents=True, exist_ok=True)

    nav_lines: List[str] = [
        "# OFFICIAL STACK - Navigation",
        "",
        "## Sections",
        "- `00_OVERVIEW`: executive summary, navigation, gate status",
        "- `01_SPECIFICATION`: formal specification and storage characteristics",
        "- `02_BENCH_METHODS`: benchmark methods and external official benchmark catalog",
        "- `03_RESULTS`: aggregated tables and Pareto outputs",
        "- `04_ANALYSIS`: advantages, memory/speed and data-impact analyses",
        "- `05_REPRO_LEGAL`: evidence manifest, source run manifests, legal checklists",
        "- `06_PRESENTATION`: xlsx/docx/pdf deliverables",
        "",
        "## Quick Status",
        f"- overall_status: `{gates.get('overall_status')}`",
        f"- run_ids: `{', '.join(str(x) for x in manifest.get('run_ids', []))}`",
        f"- tracks: `{', '.join(str(x) for x in report_json.get('tracks', []))}`",
        f"- tiers: `{', '.join(str(x) for x in report_json.get('tiers', []))}`",
        f"- scales: `{', '.join(str(x) for x in report_json.get('scales', []))}`",
        "",
    ]
    (sec_overview / "00_NAVIGATION.md").write_text("\n".join(nav_lines).strip() + "\n", encoding="utf-8")

    summary_lines: List[str] = []
    if lang in ("ru", "both"):
        summary_lines.extend([
            "# Executive Summary (RU)",
            "",
            "Пакет собран в формате официальной стопки для передачи в спецификации, due diligence и исследовательские приложения.",
            f"- gate status: `{gates.get('overall_status')}`",
            f"- run_ids: `{', '.join(str(x) for x in manifest.get('run_ids', []))}`",
            "",
        ])
    if lang in ("en", "both"):
        summary_lines.extend([
            "# Executive Summary (EN)",
            "",
            "This package is organized as an official stack suitable for specifications, due diligence, and research appendices.",
            f"- gate status: `{gates.get('overall_status')}`",
            f"- run_ids: `{', '.join(str(x) for x in manifest.get('run_ids', []))}`",
            "",
        ])
    summary_lines.append("```json")
    summary_lines.append(json.dumps(gates, ensure_ascii=False, indent=2))
    summary_lines.append("```")
    summary_lines.append("")
    (sec_overview / "01_EXECUTIVE_SUMMARY.md").write_text("\n".join(summary_lines), encoding="utf-8")
    (sec_overview / "02_ACCEPTANCE_GATES.json").write_text(
        json.dumps(gates, ensure_ascii=False, indent=2), encoding="utf-8"
    )

    _copy_to_path(artifacts["spec_md"], sec_spec / "01_SPECIFICATION.md")
    _copy_to_path(artifacts["storage_md"], sec_spec / "02_STORAGE_CHARACTERISTICS.md")

    _copy_to_path(artifacts["report_md"], sec_methods / "01_BENCH_REPORT.md")
    _copy_to_path(artifacts["external_bench_md"], sec_methods / "02_EXTERNAL_BENCHMARKS.md")
    _copy_to_path(artifacts["external_bench_csv"], sec_methods / "03_EXTERNAL_BENCHMARKS.csv")

    _copy_to_path(artifacts["summary_csv"], sec_results / "01_REPORT_SUMMARY.csv")
    _copy_to_path(artifacts["summary_json"], sec_results / "02_REPORT_SUMMARY.json")
    _copy_to_path(artifacts["evidence_matrix_csv"], sec_results / "03_EVIDENCE_MATRIX.csv")
    _copy_to_path(artifacts["memory_speed_csv"], sec_results / "04_MEMORY_SPEED.csv")
    _copy_to_path(artifacts["data_impact_csv"], sec_results / "05_DATA_IMPACT.csv")
    _copy_to_path(artifacts["pareto_latency_csv"], sec_results / "06_PARETO_RECALL_LATENCY.csv")
    _copy_to_path(artifacts["pareto_memory_csv"], sec_results / "07_PARETO_RECALL_MEMORY.csv")
    _copy_to_path(artifacts["pareto_build_csv"], sec_results / "08_PARETO_RECALL_BUILD.csv")

    _copy_to_path(artifacts["advantages_md"], sec_analysis / "01_ADVANTAGES.md")
    _copy_to_path(artifacts["memory_speed_md"], sec_analysis / "02_MEMORY_SPEED_ANALYSIS.md")
    _copy_to_path(artifacts["data_impact_md"], sec_analysis / "03_DATA_IMPACT_ANALYSIS.md")

    _copy_to_path(artifacts["evidence_manifest_json"], sec_repro / "01_EVIDENCE_MANIFEST.json")
    source_runs_dir = sec_repro / "source_runs"
    source_runs_dir.mkdir(parents=True, exist_ok=True)
    for run_dir in input_dirs:
        tag = run_dir.name
        _copy_to_path(run_dir / "manifest.json", source_runs_dir / f"{tag}_manifest.json")
        _copy_to_path(run_dir / "README_RUN.md", source_runs_dir / f"{tag}_README_RUN.md")
        _copy_to_path(run_dir / "legal_review_checklist.csv", source_runs_dir / f"{tag}_legal_review_checklist.csv")

    _copy_to_path(artifacts["xlsx"], sec_presentation / "01_DOSSIER.xlsx")
    _copy_to_path(artifacts["docx"], sec_presentation / "02_DOSSIER.docx")
    _copy_to_path(artifacts["pdf"], sec_presentation / "03_DOSSIER.pdf")
    if artifacts.get("readable_html"):
        _copy_to_path(artifacts["readable_html"], sec_presentation / "00_DOSSIER_READABLE.html")

    stack_files = _collect_files_recursive(stack_root)
    stack_manifest = {
        "generated_utc": utc_now(),
        "stack_root": str(stack_root),
        "run_ids": manifest.get("run_ids", []),
        "overall_status": gates.get("overall_status"),
        "files": [
            {
                "path": str(p.relative_to(stack_root)),
                "bytes": p.stat().st_size,
                "sha256": sha256_file(p),
            }
            for p in stack_files
        ],
    }
    (stack_root / "STACK_MANIFEST.json").write_text(
        json.dumps(stack_manifest, ensure_ascii=False, indent=2), encoding="utf-8"
    )

    stack_zip = output_dir / "OFFICIAL_STACK.zip"
    _zip_directory(stack_root, stack_zip)
    return stack_root, stack_zip


def write_excel_package(path: Path,
                        rows: Sequence[Dict[str, Any]],
                        summary_rows: Sequence[Dict[str, Any]],
                        gates: Dict[str, Any],
                        memory_speed_rows: Optional[Sequence[Dict[str, Any]]] = None,
                        data_impact_rows: Optional[Sequence[Dict[str, Any]]] = None,
                        external_benchmarks: Optional[Sequence[Dict[str, Any]]] = None) -> None:
    try:
        from openpyxl import Workbook  # type: ignore
        from openpyxl.styles import Alignment, Font, PatternFill  # type: ignore
    except Exception as exc:
        raise RuntimeError(f"openpyxl is required for xlsx export: {exc}") from exc

    def format_sheet(ws: Any, sample_rows: int = 200) -> None:
        if ws.max_row >= 1 and ws.max_column >= 1:
            ws.freeze_panes = "A2"
            ws.auto_filter.ref = ws.dimensions
            header_fill = PatternFill(fill_type="solid", start_color="1F4E78", end_color="1F4E78")
            for c in ws[1]:
                c.font = Font(bold=True, color="FFFFFF")
                c.fill = header_fill
                c.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            widths: Dict[int, int] = {}
            max_r = min(ws.max_row, sample_rows)
            for r in range(1, max_r + 1):
                for c in range(1, ws.max_column + 1):
                    v = ws.cell(r, c).value
                    ln = len(str(v)) if v is not None else 0
                    widths[c] = max(widths.get(c, 0), min(ln, 80))
            for c, w in widths.items():
                ws.column_dimensions[chr(64 + c) if c <= 26 else ws.cell(1, c).column_letter].width = max(10, min(48, w + 2))

    wb = Workbook()
    ws_info = wb.active
    ws_info.title = "README"
    ws_info.append(["Section", "Purpose"])
    ws_info.append(["README", "How to navigate this workbook"])
    ws_info.append(["kpi_overview", "Executive KPI and gate status"])
    ws_info.append(["summary", "Grouped benchmark summary"])
    ws_info.append(["memory_speed", "Throughput-memory efficiency table"])
    ws_info.append(["data_impact", "How data regimes affect performance/compactness/quality"])
    ws_info.append(["external_benchmarks", "Official external benchmark source catalog"])
    ws_info.append(["results_raw", "Full raw records, one row per measured point"])
    ws_info.append(["acceptance_gates", "Flattened gate status object"])
    format_sheet(ws_info)

    ws_kpi = wb.create_sheet("kpi_overview")
    ws_kpi.append(["KPI", "Value"])
    ws_kpi.append(["overall_status", gates.get("overall_status")])
    ws_kpi.append(["correctness_status", gates.get("gates", {}).get("correctness", {}).get("status")])
    ws_kpi.append(["reproducibility_status", gates.get("gates", {}).get("reproducibility", {}).get("status")])
    ws_kpi.append(["investor_status", gates.get("gates", {}).get("investor", {}).get("status")])
    ws_kpi.append(["scale_status", gates.get("gates", {}).get("scale", {}).get("status")])
    ws_kpi.append(["final_dossier_ready", gates.get("gates", {}).get("scale", {}).get("final_dossier_ready")])
    ws_kpi.append(["ok_scales", json.dumps(gates.get("gates", {}).get("scale", {}).get("ok_scales", []), ensure_ascii=False)])
    format_sheet(ws_kpi)

    ws_raw = wb.create_sheet("results_raw")
    ws_raw.append(list(RESULT_FIELDS))
    for row in rows:
        ws_raw.append([row.get(k) for k in RESULT_FIELDS])
    format_sheet(ws_raw, sample_rows=120)

    ws_summary = wb.create_sheet("summary")
    if summary_rows:
        headers = list(summary_rows[0].keys())
        ws_summary.append(headers)
        for row in summary_rows:
            ws_summary.append([row.get(k) for k in headers])
    else:
        ws_summary.append(["empty"])
    format_sheet(ws_summary)

    ws_gates = wb.create_sheet("acceptance_gates")
    ws_gates.append(["key", "value"])

    def flatten(prefix: str, value: Any) -> None:
        if isinstance(value, dict):
            for k, v in value.items():
                key = f"{prefix}.{k}" if prefix else str(k)
                flatten(key, v)
            return
        if isinstance(value, list):
            ws_gates.append([prefix, json.dumps(value, ensure_ascii=False)])
            return
        ws_gates.append([prefix, value])

    flatten("", gates)
    format_sheet(ws_gates, sample_rows=500)

    if memory_speed_rows is not None:
        ws_mem = wb.create_sheet("memory_speed")
        if memory_speed_rows:
            headers = list(memory_speed_rows[0].keys())
            ws_mem.append(headers)
            for row in memory_speed_rows:
                ws_mem.append([row.get(k) for k in headers])
        else:
            ws_mem.append(["empty"])
        format_sheet(ws_mem)

    if data_impact_rows is not None:
        ws_imp = wb.create_sheet("data_impact")
        if data_impact_rows:
            headers = list(data_impact_rows[0].keys())
            ws_imp.append(headers)
            for row in data_impact_rows:
                ws_imp.append([row.get(k) for k in headers])
        else:
            ws_imp.append(["empty"])
        format_sheet(ws_imp)

    if external_benchmarks is not None:
        ws_ext = wb.create_sheet("external_benchmarks")
        if external_benchmarks:
            headers = list(external_benchmarks[0].keys())
            ws_ext.append(headers)
            for row in external_benchmarks:
                ws_ext.append([row.get(k) for k in headers])
        else:
            ws_ext.append(["empty"])
        format_sheet(ws_ext)
    wb.save(path)


def write_simple_docx(path: Path, text: str) -> None:
    try:
        from docx import Document  # type: ignore

        doc = Document()
        for line in text.splitlines():
            if not line.strip():
                doc.add_paragraph("")
            elif line.startswith("# "):
                doc.add_heading(line[2:].strip(), level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:].strip(), level=2)
            else:
                doc.add_paragraph(line)
        doc.save(path)
        return
    except Exception:
        pass

    import zipfile
    from xml.sax.saxutils import escape

    paragraphs: List[str] = []
    for line in text.splitlines():
        if not line:
            paragraphs.append('<w:p/>')
            continue
        safe = escape(line)
        paragraphs.append(
            '<w:p><w:r><w:t xml:space="preserve">'
            + safe +
            "</w:t></w:r></w:p>"
        )

    document_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:document xmlns:wpc="http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas" '
        'xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006" '
        'xmlns:o="urn:schemas-microsoft-com:office:office" '
        'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" '
        'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math" '
        'xmlns:v="urn:schemas-microsoft-com:vml" '
        'xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing" '
        'xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" '
        'xmlns:w10="urn:schemas-microsoft-com:office:word" '
        'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml" '
        'xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml" '
        'xmlns:wpg="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup" '
        'xmlns:wpi="http://schemas.microsoft.com/office/word/2010/wordprocessingInk" '
        'xmlns:wne="http://schemas.microsoft.com/office/2006/wordml" '
        'xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape" mc:Ignorable="w14 w15 wp14">'
        "<w:body>"
        + "".join(paragraphs) +
        '<w:sectPr><w:pgSz w:w="11906" w:h="16838"/><w:pgMar w:top="1440" w:right="1440" w:bottom="1440" w:left="1440" '
        'w:header="708" w:footer="708" w:gutter="0"/></w:sectPr>'
        "</w:body></w:document>"
    )

    content_types = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
        '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
        '<Default Extension="xml" ContentType="application/xml"/>'
        '<Override PartName="/word/document.xml" '
        'ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
        '</Types>'
    )

    rels = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" '
        'Target="word/document.xml"/>'
        "</Relationships>"
    )

    doc_rels = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"></Relationships>'
    )

    with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("[Content_Types].xml", content_types)
        zf.writestr("_rels/.rels", rels)
        zf.writestr("word/document.xml", document_xml)
        zf.writestr("word/_rels/document.xml.rels", doc_rels)


def _fmt_num(v: Any, nd: int = 3) -> str:
    fv = to_float(v)
    if fv is None:
        return "n/a"
    return f"{fv:.{nd}f}"


def _html_table(headers: Sequence[str], rows: Sequence[Sequence[Any]]) -> str:
    out = ["<table><thead><tr>"]
    for h in headers:
        out.append(f"<th>{html.escape(str(h))}</th>")
    out.append("</tr></thead><tbody>")
    for row in rows:
        out.append("<tr>")
        for cell in row:
            out.append(f"<td>{html.escape(str(cell))}</td>")
        out.append("</tr>")
    out.append("</tbody></table>")
    return "".join(out)


def render_specialist_html_report(rows: Sequence[Dict[str, Any]],
                                  summary_rows: Sequence[Dict[str, Any]],
                                  manifest: Dict[str, Any],
                                  gates: Dict[str, Any],
                                  memory_speed_rows: Sequence[Dict[str, Any]],
                                  data_impact_rows: Sequence[Dict[str, Any]],
                                  lang: str) -> str:
    run_ids = manifest.get("run_ids") or [manifest.get("run_id")]
    tracks = sorted({str(r.get("track")) for r in rows if r.get("track") is not None})
    tiers = sorted({str(r.get("tier")) for r in rows if r.get("tier") is not None})
    scales = sorted({str(r.get("scale")) for r in rows if r.get("scale") is not None})
    ok_rows = [r for r in rows if r.get("status") == "ok"]
    top_qps = sorted(
        [r for r in memory_speed_rows if to_float(r.get("qps")) is not None],
        key=lambda x: to_float(x.get("qps")) or -1.0,
        reverse=True,
    )[:20]
    top_mem_eff = sorted(
        [r for r in memory_speed_rows if to_float(r.get("qps_per_index_mb")) is not None],
        key=lambda x: to_float(x.get("qps_per_index_mb")) or -1.0,
        reverse=True,
    )[:20]

    gate_scale = gates.get("gates", {}).get("scale", {})
    gate_rows = [
        ("overall_status", gates.get("overall_status")),
        ("correctness", gates.get("gates", {}).get("correctness", {}).get("status")),
        ("reproducibility", gates.get("gates", {}).get("reproducibility", {}).get("status")),
        ("investor", gates.get("gates", {}).get("investor", {}).get("status")),
        ("scale", gate_scale.get("status")),
        ("final_dossier_ready", gate_scale.get("final_dossier_ready")),
        ("ok_scales", ", ".join(str(x) for x in gate_scale.get("ok_scales", []))),
    ]

    summary_top = sorted(
        [r for r in summary_rows if to_float(r.get("qps")) is not None],
        key=lambda x: to_float(x.get("qps")) or -1.0,
        reverse=True,
    )[:25]

    ext_rows = [(e["category"], e["name"], e["scope"], e["url"]) for e in EXTERNAL_BENCH_LIBRARY]

    title = "Official Benchmark Dossier" if lang in ("en", "both") else "Официальный Benchmark Dossier"
    subtitle_ru = "Читаемая версия для экспертов: структура, KPI, таблицы и источники."
    subtitle_en = "Readable expert edition: structure, KPIs, analytical tables and official sources."

    css = """
    body { font-family: "Liberation Sans", Arial, sans-serif; margin: 24px; color: #1f1f1f; }
    h1 { font-size: 28px; margin: 0 0 8px 0; }
    h2 { font-size: 20px; margin-top: 28px; border-bottom: 1px solid #ddd; padding-bottom: 6px; }
    h3 { font-size: 15px; margin-top: 20px; }
    .meta { font-size: 11px; color: #444; margin-bottom: 12px; }
    .kpi-grid { display: grid; grid-template-columns: repeat(3, 1fr); gap: 10px; margin: 14px 0; }
    .kpi { border: 1px solid #d9d9d9; border-radius: 8px; padding: 10px; background: #fafafa; }
    .kpi .name { font-size: 11px; color: #555; }
    .kpi .val { font-size: 17px; font-weight: 700; margin-top: 3px; }
    table { width: 100%; border-collapse: collapse; margin: 10px 0 18px 0; font-size: 10px; }
    th, td { border: 1px solid #cfcfcf; padding: 5px; text-align: left; vertical-align: top; }
    th { background: #f0f4f8; }
    ul { margin-top: 4px; }
    .mono { font-family: "Liberation Mono", Consolas, monospace; font-size: 10px; }
    .section-note { background: #f7fbff; border-left: 3px solid #4a90e2; padding: 8px 10px; margin: 8px 0 14px 0; font-size: 11px; }
    """

    html_parts: List[str] = [
        "<!doctype html><html><head><meta charset='utf-8'>",
        f"<title>{html.escape(title)}</title>",
        f"<style>{css}</style></head><body>",
        f"<h1>{html.escape(title)}</h1>",
        f"<div>{html.escape(subtitle_ru)}</div><div>{html.escape(subtitle_en)}</div>",
        "<div class='meta'>"
        + f"generated_utc={html.escape(utc_now())} | runs={html.escape(', '.join(str(x) for x in run_ids if x))} "
        + f"| tracks={html.escape(', '.join(tracks))} | tiers={html.escape(', '.join(tiers))} | scales={html.escape(', '.join(scales))}"
        + "</div>",
        "<div class='kpi-grid'>",
        "<div class='kpi'><div class='name'>overall_status</div><div class='val'>"
        + html.escape(str(gates.get("overall_status")))
        + "</div></div>",
        "<div class='kpi'><div class='name'>rows_total</div><div class='val'>"
        + html.escape(str(len(rows)))
        + "</div></div>",
        "<div class='kpi'><div class='name'>rows_ok</div><div class='val'>"
        + html.escape(str(len(ok_rows)))
        + "</div></div>",
        "</div>",
        "<h2>1. Gate Status</h2>",
        _html_table(["Gate", "Value"], gate_rows),
        "<h2>2. Executive KPIs</h2>",
        "<div class='section-note'>Top throughput and memory-efficiency points for quick specialist review.</div>",
        "<h3>2.1 Top Throughput</h3>",
        _html_table(
            ["track", "dataset", "tier", "scale", "algorithm", "qps", "p95(ms)", "index(MB)", "qps/indexMB"],
            [
                (
                    r.get("track"),
                    r.get("dataset_id"),
                    r.get("tier"),
                    r.get("scale"),
                    r.get("algorithm"),
                    _fmt_num(r.get("qps"), 3),
                    _fmt_num(r.get("lat_p95_ms"), 6),
                    _fmt_num(r.get("index_mb"), 3),
                    _fmt_num(r.get("qps_per_index_mb"), 3),
                )
                for r in top_qps
            ],
        ),
        "<h3>2.2 Top Memory Efficiency</h3>",
        _html_table(
            ["track", "dataset", "tier", "scale", "algorithm", "qps/indexMB", "qps", "index(MB)", "recall@k"],
            [
                (
                    r.get("track"),
                    r.get("dataset_id"),
                    r.get("tier"),
                    r.get("scale"),
                    r.get("algorithm"),
                    _fmt_num(r.get("qps_per_index_mb"), 3),
                    _fmt_num(r.get("qps"), 3),
                    _fmt_num(r.get("index_mb"), 3),
                    _fmt_num(r.get("recall_at_k"), 3),
                )
                for r in top_mem_eff
            ],
        ),
        "<h2>3. Data Impact Analysis</h2>",
        _html_table(
            ["track", "dataset", "regime", "tier", "scale", "mean_qps", "mean_p95(ms)", "mean_bytes/insert", "mean_recall"],
            [
                (
                    r.get("track"),
                    r.get("dataset_id"),
                    r.get("data_regime"),
                    r.get("tier"),
                    r.get("scale"),
                    _fmt_num(r.get("mean_qps"), 3),
                    _fmt_num(r.get("mean_lat_p95_ms"), 6),
                    _fmt_num(r.get("mean_bytes_per_insert"), 3),
                    _fmt_num(r.get("mean_recall_at_k"), 3),
                )
                for r in data_impact_rows[:120]
            ],
        ),
        "<h2>4. Summary Table (Top by QPS)</h2>",
        _html_table(
            ["dataset", "tier", "scale", "algorithm", "track", "status", "qps", "p95(ms)", "recall@k", "bytes/insert"],
            [
                (
                    r.get("dataset_id"),
                    r.get("tier"),
                    r.get("scale"),
                    r.get("algorithm"),
                    r.get("track"),
                    r.get("status"),
                    _fmt_num(r.get("qps"), 3),
                    _fmt_num(r.get("lat_p95_ms"), 6),
                    _fmt_num(r.get("recall_at_k"), 3),
                    _fmt_num(r.get("bytes_per_insert"), 3),
                )
                for r in summary_top
            ],
        ),
        "<h2>5. External Official Benchmarks</h2>",
        _html_table(["category", "benchmark", "scope", "url"], ext_rows),
        "<h2>6. Notes</h2>",
        "<ul>"
        "<li>This readable dossier is generated from machine-validated CSV/JSON artifacts in the same output folder.</li>"
        "<li>For reproducibility/audit, use <span class='mono'>OFFICIAL_STACK/05_REPRO_LEGAL</span>.</li>"
        "<li>For full raw points, use <span class='mono'>report_summary.csv</span> and <span class='mono'>dossier_evidence_matrix.csv</span>.</li>"
        "</ul>",
        "</body></html>",
    ]
    return "".join(html_parts)


def write_html_file(path: Path, html_text: str) -> None:
    path.write_text(html_text, encoding="utf-8")


def convert_html_with_soffice(html_path: Path, out_path: Path, convert_spec: str) -> bool:
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return False
    out_path.parent.mkdir(parents=True, exist_ok=True)
    cp = safe_run([soffice, "--headless", "--convert-to", convert_spec, "--outdir", str(out_path.parent), str(html_path)])
    if cp.returncode != 0:
        return False
    produced = html_path.with_suffix("." + out_path.suffix.lstrip("."))
    if not produced.exists():
        produced = out_path.parent / (html_path.stem + "." + out_path.suffix.lstrip("."))
    if not produced.exists():
        return False
    if produced != out_path:
        shutil.move(str(produced), str(out_path))
    return out_path.exists()


def render_markdown_report(rows: Sequence[Dict[str, Any]], lang: str, manifest: Dict[str, Any]) -> str:
    def fmt(v: Any, nd: int = 3) -> str:
        if v is None:
            return "—"
        if isinstance(v, float):
            return f"{v:.{nd}f}"
        return str(v)

    by_track: Dict[str, List[Dict[str, Any]]] = {}
    for r in rows:
        by_track.setdefault(str(r.get("track", "unknown")), []).append(r)
    run_ids = manifest.get("run_ids")
    if isinstance(run_ids, list) and run_ids:
        run_label = ", ".join(str(x) for x in run_ids)
    else:
        run_label = str(manifest.get("run_id"))

    lines: List[str] = []
    if lang in ("ru", "both"):
        lines.append("# Benchmark Zone Report (RU)")
        lines.append("")
        lines.append("## Контур")
        lines.append(f"- run_id(s): `{run_label}`")
        lines.append(f"- created_utc: `{manifest.get('created_utc')}`")
        lines.append(f"- git_commit: `{manifest.get('git_commit')}`")
        lines.append(f"- licensing_policy: `{manifest.get('licensing_policy')}`")
        lines.append("")
        lines.append("## Ключевые итоги")
        for track, tr_rows in by_track.items():
            ok = sum(1 for r in tr_rows if r.get("status") == "ok")
            skip = sum(1 for r in tr_rows if r.get("status") == "skipped")
            fail = sum(1 for r in tr_rows if r.get("status") == "failed")
            qps_mean = numeric_mean([r.get("qps") for r in tr_rows if isinstance(r.get("qps"), (int, float))])
            lines.append(f"- `{track}`: ok={ok}, skipped={skip}, failed={fail}, mean_qps={fmt(qps_mean)}")
        lines.append("")

    if lang in ("en", "both"):
        lines.append("# Benchmark Zone Report (EN)")
        lines.append("")
        lines.append("## Context")
        lines.append(f"- run_id(s): `{run_label}`")
        lines.append(f"- created_utc: `{manifest.get('created_utc')}`")
        lines.append(f"- git_commit: `{manifest.get('git_commit')}`")
        lines.append(f"- licensing_policy: `{manifest.get('licensing_policy')}`")
        lines.append("")
        lines.append("## Key outcomes")
        for track, tr_rows in by_track.items():
            ok = sum(1 for r in tr_rows if r.get("status") == "ok")
            skip = sum(1 for r in tr_rows if r.get("status") == "skipped")
            fail = sum(1 for r in tr_rows if r.get("status") == "failed")
            qps_mean = numeric_mean([r.get("qps") for r in tr_rows if isinstance(r.get("qps"), (int, float))])
            lines.append(f"- `{track}`: ok={ok}, skipped={skip}, failed={fail}, mean_qps={fmt(qps_mean)}")
        lines.append("")

    lines.append("## Top rows (sample)")
    lines.append("")
    lines.append("| track | dataset | algorithm | tier | scale | qps | p95 ms | recall@k | min_hamming | audit_gap | status |")
    lines.append("|---|---|---|---|---|---:|---:|---:|---:|---:|---|")
    preview = list(rows)[:30]
    for r in preview:
        lines.append(
            "| {track} | {dataset_id} | {algorithm} | {tier} | {scale} | {qps} | {p95} | {recall} | {hamm} | {gap} | {status} |".format(
                track=r.get("track"),
                dataset_id=r.get("dataset_id"),
                algorithm=r.get("algorithm"),
                tier=r.get("tier"),
                scale=r.get("scale"),
                qps=fmt(r.get("qps")),
                p95=fmt(r.get("lat_p95_ms")),
                recall=fmt(r.get("recall_at_k")),
                hamm=fmt(r.get("error_bits_min_hamming")),
                gap=fmt(r.get("hamming_audit_gap")),
                status=r.get("status"),
            )
        )
    lines.append("")
    lines.append("## Reproducibility")
    lines.append("- `results.jsonl` contains raw run-level records.")
    lines.append("- `summary.csv` contains grouped mean metrics.")
    lines.append("- `manifest.json` contains hardware/software fingerprints and source pack.")
    lines.append("- `legal_review_checklist.csv` tracks publication gate.")
    lines.append("")
    return "\n".join(lines).strip() + "\n"


def write_simple_pdf(path: Path, text: str) -> None:
    lines = text.splitlines()
    max_lines = 90
    if len(lines) > max_lines:
        lines = lines[: max_lines - 2] + ["...", "[truncated for single-page PDF exporter]"]

    escaped = []
    for line in lines:
        s = line.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
        escaped.append(s)

    content_parts = ["BT", "/F1 9 Tf", "48 800 Td"]
    first = True
    for line in escaped:
        if first:
            content_parts.append(f"({line}) Tj")
            first = False
        else:
            content_parts.append("0 -10 Td")
            content_parts.append(f"({line}) Tj")
    content_parts.append("ET")
    stream = ("\n".join(content_parts) + "\n").encode("latin-1", errors="replace")

    objs: List[bytes] = []
    objs.append(b"<< /Type /Catalog /Pages 2 0 R >>")
    objs.append(b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>")
    objs.append(b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 595 842] /Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>")
    objs.append(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")
    objs.append(b"<< /Length " + str(len(stream)).encode("ascii") + b" >>\nstream\n" + stream + b"endstream")

    chunks: List[bytes] = [b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n"]
    offsets = [0]
    for i, obj in enumerate(objs, start=1):
        offsets.append(sum(len(c) for c in chunks))
        chunks.append(f"{i} 0 obj\n".encode("ascii"))
        chunks.append(obj)
        chunks.append(b"\nendobj\n")
    xref_pos = sum(len(c) for c in chunks)
    chunks.append(f"xref\n0 {len(objs)+1}\n".encode("ascii"))
    chunks.append(b"0000000000 65535 f \n")
    for off in offsets[1:]:
        chunks.append(f"{off:010d} 00000 n \n".encode("ascii"))
    chunks.append(
        f"trailer\n<< /Size {len(objs)+1} /Root 1 0 R >>\nstartxref\n{xref_pos}\n%%EOF\n".encode("ascii")
    )
    path.write_bytes(b"".join(chunks))


def run_command(args: argparse.Namespace) -> int:
    repo_root = Path.cwd()
    cfg_path = Path(args.config)
    cfg = load_config(cfg_path)
    compile_core_tool(repo_root, cfg)

    tracks = [args.track] if args.track else [t for t in TRACK_CHOICES if cfg["tracks"].get(t, {}).get("enabled", True)]
    tiers = [args.tier] if args.tier else list(TIER_CHOICES)
    scales = [args.scale] if args.scale else list(SCALE_CHOICES)

    run_id = dt.datetime.now(dt.timezone.utc).strftime("%Y%m%dT%H%M%SZ") + "_" + uuid.uuid4().hex[:8]
    output_root = Path(cfg.get("output_root", "bench_runs"))
    run_dir = prepare_run_dir(output_root, run_id)
    seed = int(cfg.get("seed", 123456789))

    rows = run_tracks(repo_root, cfg, run_id, tracks, tiers, scales, seed)
    manifest = make_manifest(repo_root, cfg, run_id, cfg_path)
    container_hash = manifest.get("container_hash")
    hw_hash = manifest.get("hardware_fingerprint_hash")
    for r in rows:
        r["container_hash"] = container_hash
        r["hardware_fingerprint"] = hw_hash
        r["legal_review_required"] = True

    write_jsonl(run_dir / "results.jsonl", rows)
    write_summary_csv(run_dir / "summary.csv", rows)
    (run_dir / "manifest.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    write_legal_checklist(run_dir, cfg)

    readme = textwrap.dedent(
        f"""\
        # Benchmark Run {run_id}

        - created_utc: {manifest["created_utc"]}
        - tracks: {", ".join(tracks)}
        - tiers: {", ".join(tiers)}
        - scales: {", ".join(scales)}
        - results: `results.jsonl`
        - summary: `summary.csv`
        - manifest: `manifest.json`
        - legal gate: `legal_review_checklist.csv`
        """
    ).strip() + "\n"
    (run_dir / "README_RUN.md").write_text(readme, encoding="utf-8")

    print(str(run_dir))
    return 0


def report_command(args: argparse.Namespace) -> int:
    input_dirs = [Path(x) for x in args.input]
    if not input_dirs:
        raise ValueError("at least one --input directory is required")

    rows: List[Dict[str, Any]] = []
    manifests: List[Dict[str, Any]] = []
    for run_dir in input_dirs:
        results_path = run_dir / "results.jsonl"
        manifest_path = run_dir / "manifest.json"
        if not results_path.exists():
            raise FileNotFoundError(f"missing results file: {results_path}")
        if not manifest_path.exists():
            raise FileNotFoundError(f"missing manifest file: {manifest_path}")
        with results_path.open("r", encoding="utf-8") as f:
            for line in f:
                s = line.strip()
                if not s:
                    continue
                rows.append(json.loads(s))
        manifests.append(json.loads(manifest_path.read_text(encoding="utf-8")))

    if args.output:
        output_dir = Path(args.output)
    elif len(input_dirs) == 1:
        output_dir = input_dirs[0]
    else:
        output_dir = input_dirs[0].parent / (
            "dossier_" + dt.datetime.now(dt.timezone.utc).strftime("%Y%m%dT%H%M%SZ") + "_" + uuid.uuid4().hex[:8]
        )
    output_dir.mkdir(parents=True, exist_ok=True)

    primary = manifests[0]
    manifest: Dict[str, Any] = {
        "run_id": primary.get("run_id"),
        "run_ids": [m.get("run_id") for m in manifests],
        "created_utc": utc_now(),
        "git_commit": primary.get("git_commit"),
        "licensing_policy": primary.get("licensing_policy"),
        "official_source_pack": OFFICIAL_SOURCE_PACK,
        "source_run_dirs": [str(p) for p in input_dirs],
    }

    fmt = args.format
    lang = args.lang
    summary_rows = build_summary_rows(rows)
    gates = compute_acceptance_gates(rows)

    report_json = {
        "run_ids": manifest.get("run_ids"),
        "generated_utc": utc_now(),
        "n_rows": len(rows),
        "tracks": sorted(set(str(r.get("track")) for r in rows)),
        "tiers": sorted(set(str(r.get("tier")) for r in rows)),
        "scales": sorted(set(str(r.get("scale")) for r in rows)),
        "acceptance_gates": gates,
    }

    if fmt in ("csv", "all", "xlsx"):
        write_summary_csv(output_dir / "report_summary.csv", rows)
    if fmt in ("json", "all"):
        (output_dir / "report_summary.json").write_text(
            json.dumps(report_json, ensure_ascii=False, indent=2), encoding="utf-8"
        )

    report_md = render_markdown_report(rows, lang, manifest)
    report_md_path = output_dir / ("report_bilingual.md" if lang == "both" else f"report_{lang}.md")
    if fmt in ("md", "pdf", "docx", "all"):
        report_md_path.write_text(report_md, encoding="utf-8")

    memory_speed_rows = build_memory_speed_rows(rows)
    data_impact_rows = build_data_impact_rows(rows)
    specialist_html_text = render_specialist_html_report(
        rows=rows,
        summary_rows=summary_rows,
        manifest=manifest,
        gates=gates,
        memory_speed_rows=memory_speed_rows,
        data_impact_rows=data_impact_rows,
        lang=lang,
    )
    specialist_html_path = output_dir / (
        "dossier_readable_bilingual.html" if lang == "both" else f"dossier_readable_{lang}.html"
    )
    if fmt in ("md", "pdf", "docx", "all"):
        write_html_file(specialist_html_path, specialist_html_text)

    spec_text = (
        render_specification_doc("ru", manifest) + "\n" + render_specification_doc("en", manifest)
        if lang == "both"
        else render_specification_doc(lang, manifest)
    )
    adv_text = (
        render_advantages_doc(rows, "ru") + "\n" + render_advantages_doc(rows, "en")
        if lang == "both"
        else render_advantages_doc(rows, lang)
    )
    memory_speed_text = render_memory_speed_report(rows, lang)
    data_impact_text = render_data_impact_report(rows, lang)
    storage_text = render_storage_characteristics_report(rows, lang)
    external_bench_text = render_external_benchmarks_report(lang)

    spec_md_path = output_dir / ("dossier_specification_bilingual.md" if lang == "both" else f"dossier_specification_{lang}.md")
    adv_md_path = output_dir / ("dossier_advantages_bilingual.md" if lang == "both" else f"dossier_advantages_{lang}.md")
    memory_speed_md_path = output_dir / ("dossier_memory_speed_bilingual.md" if lang == "both" else f"dossier_memory_speed_{lang}.md")
    data_impact_md_path = output_dir / ("dossier_data_impact_bilingual.md" if lang == "both" else f"dossier_data_impact_{lang}.md")
    storage_md_path = output_dir / ("dossier_storage_characteristics_bilingual.md" if lang == "both" else f"dossier_storage_characteristics_{lang}.md")
    external_bench_md_path = output_dir / ("dossier_external_benchmarks_bilingual.md" if lang == "both" else f"dossier_external_benchmarks_{lang}.md")
    spec_md_path.write_text(spec_text, encoding="utf-8")
    adv_md_path.write_text(adv_text, encoding="utf-8")
    memory_speed_md_path.write_text(memory_speed_text, encoding="utf-8")
    data_impact_md_path.write_text(data_impact_text, encoding="utf-8")
    storage_md_path.write_text(storage_text, encoding="utf-8")
    external_bench_md_path.write_text(external_bench_text, encoding="utf-8")
    write_csv_rows(
        output_dir / "dossier_memory_speed_report.csv",
        memory_speed_rows,
        fieldnames=(
            "track",
            "dataset_id",
            "tier",
            "scale",
            "algorithm",
            "status",
            "qps",
            "lat_p95_ms",
            "index_mb",
            "rss_mb",
            "bytes_per_insert",
            "recall_at_k",
            "error_bits_min_hamming",
            "qps_per_index_mb",
            "qps_per_rss_mb",
        ),
    )
    write_csv_rows(
        output_dir / "dossier_data_impact_report.csv",
        data_impact_rows,
        fieldnames=(
            "track",
            "dataset_id",
            "data_regime",
            "tier",
            "scale",
            "rows_in_group",
            "mean_qps",
            "mean_lat_p95_ms",
            "mean_bytes_per_insert",
            "mean_recall_at_k",
            "mean_error_bits_min_hamming",
            "best_algorithm_by_qps",
            "best_qps",
            "best_algorithm_by_compactness",
            "best_qps_per_index_mb",
        ),
    )
    write_external_benchmark_library_csv(output_dir / "dossier_external_benchmarks.csv")

    dossier_text = (
        report_md
        + "\n## Acceptance Gates\n\n```json\n"
        + json.dumps(gates, ensure_ascii=False, indent=2)
        + "\n```\n\n"
        + spec_text
        + "\n"
        + adv_text
        + "\n"
        + memory_speed_text
        + "\n"
        + data_impact_text
        + "\n"
        + storage_text
        + "\n"
        + external_bench_text
    )
    dossier_md_path = output_dir / ("dossier_bilingual.md" if lang == "both" else f"dossier_{lang}.md")
    dossier_md_path.write_text(dossier_text, encoding="utf-8")

    (output_dir / "dossier_acceptance_gates.json").write_text(
        json.dumps(gates, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    write_evidence_matrix_csv(output_dir / "dossier_evidence_matrix.csv", rows)
    write_ann_pareto_csv(
        output_dir / "dossier_pareto_recall_latency.csv",
        rows,
        maximize_fields=("recall_at_k",),
        minimize_fields=("lat_p95_ms",),
        value_fields=("recall_at_k", "lat_p95_ms", "qps", "index_bytes", "build_time_s"),
    )
    write_ann_pareto_csv(
        output_dir / "dossier_pareto_recall_memory.csv",
        rows,
        maximize_fields=("recall_at_k",),
        minimize_fields=("index_bytes",),
        value_fields=("recall_at_k", "index_bytes", "rss_mb", "qps", "build_time_s"),
    )
    write_ann_pareto_csv(
        output_dir / "dossier_pareto_recall_build.csv",
        rows,
        maximize_fields=("recall_at_k",),
        minimize_fields=("build_time_s",),
        value_fields=("recall_at_k", "build_time_s", "qps", "index_bytes", "lat_p95_ms"),
    )

    pdf_path = output_dir / ("report_bilingual.pdf" if lang == "both" else f"report_{lang}.pdf")
    dossier_pdf = output_dir / ("dossier_bilingual.pdf" if lang == "both" else f"dossier_{lang}.pdf")
    if fmt in ("pdf", "all"):
        ok_report_pdf = convert_html_with_soffice(specialist_html_path, pdf_path, "pdf:writer_web_pdf_Export")
        if not ok_report_pdf:
            write_simple_pdf(pdf_path, report_md)
        ok_dossier_pdf = convert_html_with_soffice(specialist_html_path, dossier_pdf, "pdf:writer_web_pdf_Export")
        if not ok_dossier_pdf:
            write_simple_pdf(dossier_pdf, dossier_text)

    if fmt in ("xlsx", "all"):
        write_excel_package(
            output_dir / "dossier_package.xlsx",
            rows,
            summary_rows,
            gates,
            memory_speed_rows=memory_speed_rows,
            data_impact_rows=data_impact_rows,
            external_benchmarks=[dict(x) for x in EXTERNAL_BENCH_LIBRARY],
        )

    if fmt in ("docx", "all"):
        docx_out = output_dir / "dossier_package.docx"
        ok_docx = convert_html_with_soffice(specialist_html_path, docx_out, "docx:MS Word 2007 XML")
        if not ok_docx:
            write_simple_docx(docx_out, dossier_text)

    artifacts = {
        "report_md": report_md_path,
        "summary_csv": output_dir / "report_summary.csv",
        "summary_json": output_dir / "report_summary.json",
        "spec_md": spec_md_path,
        "advantages_md": adv_md_path,
        "memory_speed_md": memory_speed_md_path,
        "data_impact_md": data_impact_md_path,
        "storage_md": storage_md_path,
        "external_bench_md": external_bench_md_path,
        "evidence_matrix_csv": output_dir / "dossier_evidence_matrix.csv",
        "memory_speed_csv": output_dir / "dossier_memory_speed_report.csv",
        "data_impact_csv": output_dir / "dossier_data_impact_report.csv",
        "external_bench_csv": output_dir / "dossier_external_benchmarks.csv",
        "pareto_latency_csv": output_dir / "dossier_pareto_recall_latency.csv",
        "pareto_memory_csv": output_dir / "dossier_pareto_recall_memory.csv",
        "pareto_build_csv": output_dir / "dossier_pareto_recall_build.csv",
        "evidence_manifest_json": output_dir / "dossier_evidence_manifest.json",
        "readable_html": specialist_html_path,
        "xlsx": output_dir / "dossier_package.xlsx",
        "docx": output_dir / "dossier_package.docx",
        "pdf": dossier_pdf,
    }

    evidence_files = [
        output_dir / "report_summary.csv",
        output_dir / "report_summary.json",
        report_md_path,
        spec_md_path,
        adv_md_path,
        memory_speed_md_path,
        data_impact_md_path,
        storage_md_path,
        external_bench_md_path,
        dossier_md_path,
        output_dir / "dossier_acceptance_gates.json",
        output_dir / "dossier_evidence_matrix.csv",
        output_dir / "dossier_memory_speed_report.csv",
        output_dir / "dossier_data_impact_report.csv",
        output_dir / "dossier_external_benchmarks.csv",
        output_dir / "dossier_pareto_recall_latency.csv",
        output_dir / "dossier_pareto_recall_memory.csv",
        output_dir / "dossier_pareto_recall_build.csv",
        specialist_html_path,
        output_dir / "dossier_package.xlsx",
        output_dir / "dossier_package.docx",
        output_dir / "dossier_bilingual.pdf",
        output_dir / "dossier_en.pdf",
        output_dir / "dossier_ru.pdf",
    ]
    write_evidence_manifest(
        output_dir / "dossier_evidence_manifest.json",
        evidence_files,
        {"input_dirs": [str(p) for p in input_dirs], "report_json": report_json},
    )

    stack_root = output_dir / "OFFICIAL_STACK"
    stack_zip = output_dir / "OFFICIAL_STACK.zip"
    build_official_stack_bundle(
        output_dir=output_dir,
        input_dirs=input_dirs,
        lang=lang,
        manifest=manifest,
        report_json=report_json,
        gates=gates,
        artifacts=artifacts,
    )

    # Refresh evidence manifest including OFFICIAL_STACK bundle and final checksums.
    write_evidence_manifest(
        output_dir / "dossier_evidence_manifest.json",
        evidence_files + [stack_zip, stack_root / "STACK_MANIFEST.json", stack_root / "00_OVERVIEW" / "00_NAVIGATION.md"],
        {"input_dirs": [str(p) for p in input_dirs], "report_json": report_json, "official_stack": str(stack_root)},
    )

    print(str(output_dir))
    return 0


def build_arg_parser() -> argparse.ArgumentParser:
    ap = argparse.ArgumentParser(description="Large benchmark zone orchestrator for core256/simplefill")
    sub = ap.add_subparsers(dest="cmd", required=True)

    run = sub.add_parser("run", help="run benchmark tracks")
    run.add_argument("--config", default="bench_zone.yaml", help="path to benchmark yaml config")
    run.add_argument("--track", choices=TRACK_CHOICES, default=None, help="single track override")
    run.add_argument("--tier", choices=TIER_CHOICES, default=None, help="single tier override")
    run.add_argument("--scale", choices=SCALE_CHOICES, default=None, help="single scale override")
    run.set_defaults(func=run_command)

    rep = sub.add_parser("report", help="build report artifacts from run directory")
    rep.add_argument("--input", required=True, nargs="+", help="one or more benchmark run directories")
    rep.add_argument("--output", default=None, help="output directory for dossier artifacts")
    rep.add_argument("--lang", choices=("ru", "en", "both"), default="both")
    rep.add_argument("--format", choices=("md", "pdf", "csv", "json", "xlsx", "docx", "all"), default="all")
    rep.set_defaults(func=report_command)
    return ap


def main(argv: Optional[Sequence[str]] = None) -> int:
    parser = build_arg_parser()
    args = parser.parse_args(argv)
    try:
        return int(args.func(args))
    except Exception as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
